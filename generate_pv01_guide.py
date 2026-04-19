# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/workspace/phong_van_AI_phap_che"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), 'AAAAAA')
        tcBorders = OxmlElement('w:tcBorders')
        tcBorders.append(border)
        tcPr.append(tcBorders)


def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def add_h3(doc, text, color=(0x37, 0x86, 0x4E)):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)


def add_body(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.italic = italic
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_node_badge(doc, nodes):
    """Tạo dòng hiển thị các node badge"""
    p = doc.add_paragraph()
    p.add_run("→ NODE: ").bold = True
    for i, (node_code, node_name, color) in enumerate(nodes):
        run = p.add_run(f"  [{node_code}] {node_name}  ")
        run.bold = True
        run.font.color.rgb = RGBColor(*color)
        if i < len(nodes) - 1:
            p.add_run("  +  ")
    return p


def add_reason(doc, text):
    p = doc.add_paragraph()
    p.add_run("Lý do: ").bold = True
    r = p.add_run(text)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return p


def add_warning(doc, text):
    p = doc.add_paragraph()
    r1 = p.add_run("⚠ Lưu ý: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
    return p


def add_memo(doc, text):
    p = doc.add_paragraph()
    r1 = p.add_run("✎ Gợi ý Memo: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x5B, 0x2C, 0x8D)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x5B, 0x2C, 0x8D)
    r2.font.italic = True
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)


def add_data_box(doc, label, text, bg_color, label_color):
    """Tạo ô dữ liệu có màu nền"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label}\n")
    r1.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()
    return table


# Màu sắc cho từng nhóm node
COLOR_BOI_CANH   = (0x1F, 0x4E, 0x79)   # xanh đậm
COLOR_UNG_DUNG   = (0x37, 0x86, 0x4E)   # xanh lá
COLOR_LOI_ICH    = (0xBF, 0x85, 0x00)   # vàng nâu
COLOR_THACH_THUC = (0xC0, 0x50, 0x20)   # cam đỏ
COLOR_PHAP_LY    = (0x5B, 0x2C, 0x8D)   # tím
COLOR_DAO_DUC    = (0x96, 0x23, 0x1C)   # đỏ
COLOR_TUONG_LAI  = (0x40, 0x40, 0x40)   # xám đen


def create_pv01_coding_guide():
    doc = Document()

    # ── TRANG BÌA ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_title(doc, "HƯỚNG DẪN CODE SELECTION CHI TIẾT")
    p = doc.add_paragraph("File phỏng vấn: PV01 – Ông Nguyễn Minh Trí")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p2 = doc.add_paragraph("Trưởng phòng Pháp chế – Công ty CP Tài chính Số Việt (FinTech)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("Mục đích", "Hướng dẫn từng bước bôi chọn (code selection) toàn bộ PV01 trong NVivo"),
        ("Số câu hỏi", "8 câu hỏi / 8 câu trả lời của NTL"),
        ("Tổng số đoạn cần mã hóa", "21 đoạn mã hóa được đề xuất"),
        ("Codebook sử dụng", "CODEBOOK_Ma_Hoa_AI_Phap_Che.docx – phiên bản 1.0"),
    ]
    for i, (k, v) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = k
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(row[0], "D6E4F7")
        row[1].text = v

    doc.add_paragraph()

    # Bảng màu node legend
    add_h2(doc, "KÝ HIỆU MÀU SẮC NODE")
    legend_table = doc.add_table(rows=1, cols=7)
    legend_table.style = 'Table Grid'
    legend_data = [
        ("BCLC\nBối cảnh", "1F4E79"),
        ("UNG_DUNG\nỨng dụng", "37864E"),
        ("LOI_ICH\nLợi ích", "BF8500"),
        ("THACH_THUC\nThách thức", "C05020"),
        ("PHAP_LY\nPháp lý", "5B2C8D"),
        ("DAO_DUC\nĐạo đức", "96231C"),
        ("TUONG_LAI\nTương lai", "404040"),
    ]
    for i, (label, hex_col) in enumerate(legend_data):
        cell = legend_table.rows[0].cells[i]
        cell.text = label
        r, g, b = int(hex_col[0:2], 16), int(hex_col[2:4], 16), int(hex_col[4:6], 16)
        # Làm màu nền nhạt hơn
        r2 = min(255, r + 130)
        g2 = min(255, g + 130)
        b2 = min(255, b + 130)
        light = f"{r2:02X}{g2:02X}{b2:02X}"
        set_cell_bg(cell, light)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── CÁCH ĐỌC FILE NÀY ─────────────────────────────────────────────────
    add_h2(doc, "CÁCH SỬ DỤNG FILE HƯỚNG DẪN NÀY")
    steps = [
        "Mở NVivo → Mở file PV01 ở một màn hình / cửa sổ.",
        "Mở file hướng dẫn này ở màn hình còn lại (hoặc in ra giấy).",
        "Đọc từng MỤC bên dưới theo thứ tự câu hỏi Q1 → Q8.",
        "Với mỗi mục: đọc ĐOẠN CẦN MÃ HÓA → nhìn vào NODE ĐỀ XUẤT → bôi chọn đúng đoạn đó trong NVivo → gán node.",
        "Một đoạn có nhiều node → bôi lại và gán từng node một.",
        "Đọc LÝ DO để hiểu tại sao chọn node đó.",
        "Đọc ⚠ LƯU Ý nếu có – những chỗ dễ nhầm lẫn.",
        "Đọc ✎ GỢI Ý MEMO nếu có – những suy nghĩ đáng ghi lại.",
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')
    doc.add_paragraph()

    add_body(doc,
        "Lưu ý quan trọng: Các đề xuất dưới đây là GỢI Ý, không phải quy tắc cứng. "
        "Bạn hoàn toàn có thể không đồng ý và chọn node khác – miễn là ghi lý do vào Memo.",
        italic=True
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # Q1
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q1 – Giới thiệu bản thân và bộ phận pháp chế")
    add_body(doc, "Câu hỏi PVV: Ông có thể giới thiệu sơ lược về vai trò của mình tại công ty và bộ phận pháp chế không?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 1.1", color=COLOR_BOI_CANH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Tôi làm Trưởng phòng Pháp chế tại Công ty CP Tài chính Số Việt được gần 5 năm nay. "
        "Công ty chúng tôi hoạt động trong lĩnh vực fintech – cho vay ngang hàng, ví điện tử và các dịch vụ tài chính số.",
        "EBF3FC", RGBColor(0x1F, 0x4E, 0x79)
    )
    add_node_badge(doc, [
        ("BCLC.04", "Đặc thù ngành / lĩnh vực", COLOR_BOI_CANH),
    ])
    add_reason(doc, "Đoạn này xác định đặc thù của ngành fintech – bối cảnh quan trọng để hiểu tại sao AI cần thiết. Không mã hóa thông tin cá nhân (tên, chức vụ) vì đó chỉ là metadata.")
    add_warning(doc, "Không mã hóa toàn bộ câu trả lời Q1 vào một node. Chỉ chọn phần nói về đặc thù ngành.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 1.2", color=COLOR_BOI_CANH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Đặc thù của fintech là môi trường pháp lý thay đổi rất nhanh, nên khối lượng công việc luôn rất lớn.",
        "EBF3FC", RGBColor(0x1F, 0x4E, 0x79)
    )
    add_node_badge(doc, [
        ("BCLC.01", "Áp lực & quá tải công việc", COLOR_BOI_CANH),
        ("BCLC.04", "Đặc thù ngành / lĩnh vực", COLOR_BOI_CANH),
    ])
    add_reason(doc, "Một câu ngắn nhưng chứa hai ý: (1) đặc thù fintech thay đổi nhanh → BCLC.04; (2) khối lượng công việc lớn → BCLC.01. Đây là ví dụ điển hình của multiple coding.")
    add_memo(doc, "Ghi chú: NTL dùng cụm 'môi trường pháp lý thay đổi rất nhanh' như lý do cốt lõi – cần theo dõi xem các NTL khác có cùng nhận định về tốc độ thay đổi không.")

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q2
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q2 – Bắt đầu ứng dụng AI từ khi nào và tại sao")
    add_body(doc, "Câu hỏi PVV: Công ty đã bắt đầu ứng dụng AI vào công tác pháp chế từ khi nào và xuất phát từ nhu cầu gì?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 2.1", color=COLOR_BOI_CANH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Chúng tôi bắt đầu thử nghiệm từ cuối năm 2022, nhưng triển khai chính thức vào đầu năm 2023.",
        "EBF3FC", RGBColor(0x1F, 0x4E, 0x79)
    )
    add_node_badge(doc, [
        ("BCLC.03", "Lộ trình & giai đoạn triển khai", COLOR_BOI_CANH),
    ])
    add_reason(doc, "Câu này xác định mốc thời gian rõ ràng: thử nghiệm 2022, chính thức 2023 – đây là thông tin về lộ trình triển khai.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 2.2", color=COLOR_BOI_CANH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Nhu cầu xuất phát từ thực tế là bộ phận pháp chế của chúng tôi liên tục bị quá tải. "
        "Mỗi tháng chúng tôi phải rà soát hàng trăm hợp đồng đối tác, hàng chục điều khoản điều kiện sử dụng dịch vụ, "
        "chưa kể việc theo dõi văn bản pháp luật mới.",
        "EBF3FC", RGBColor(0x1F, 0x4E, 0x79)
    )
    add_node_badge(doc, [
        ("BCLC.01", "Áp lực & quá tải công việc", COLOR_BOI_CANH),
    ])
    add_reason(doc, "Đây là đoạn mô tả rõ nhất về áp lực công việc với số liệu cụ thể (hàng trăm hợp đồng/tháng). Đây là reference quan trọng nhất cho BCLC.01 trong toàn bộ PV01.")
    add_memo(doc, "Số liệu định lượng hiếm gặp trong phỏng vấn định tính – ghi chú để có thể dùng làm evidence trong báo cáo.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 2.3", color=COLOR_BOI_CANH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Ban đầu chúng tôi dùng các công cụ AI đơn giản để tóm tắt văn bản. Sau đó chúng tôi đầu tư vào một nền tảng AI pháp lý "
        "chuyên biệt có khả năng phân tích hợp đồng, nhận diện rủi ro, và so sánh điều khoản với mẫu chuẩn của chúng tôi.",
        "EBF3FC", RGBColor(0x1F, 0x4E, 0x79)
    )
    add_node_badge(doc, [
        ("BCLC.03", "Lộ trình & giai đoạn triển khai", COLOR_BOI_CANH),
        ("UNG_DUNG.01", "Rà soát & phân tích hợp đồng", COLOR_UNG_DUNG),
    ])
    add_reason(doc,
        "Câu đầu nói về quá trình phát triển dần dần (lộ trình) → BCLC.03. "
        "Câu sau mô tả tính năng cụ thể của AI: phân tích hợp đồng, nhận diện rủi ro → UNG_DUNG.01."
    )
    add_warning(doc, "Đây là đoạn multiple coding. Bôi chọn toàn bộ đoạn → gán BCLC.03, rồi bôi lại câu thứ hai → gán UNG_DUNG.01.")

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q3
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q3 – AI hỗ trợ công việc hàng ngày như thế nào")
    add_body(doc, "Câu hỏi PVV: Cụ thể AI đang hỗ trợ bộ phận của ông như thế nào trong công việc hàng ngày?", italic=True)
    doc.add_paragraph()
    add_body(doc, "★ Đây là câu trả lời quan trọng nhất trong PV01, chứa nhiều ứng dụng AI cụ thể nhất. Cần mã hóa cẩn thận.", bold=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 3.1", color=COLOR_UNG_DUNG)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thứ nhất là rà soát hợp đồng tự động – AI quét toàn bộ hợp đồng, đánh dấu các điều khoản bất lợi, "
        "thiếu sót hoặc mơ hồ, và so sánh với thư viện điều khoản chuẩn của chúng tôi.",
        "EBF5EA", RGBColor(0x37, 0x86, 0x4E)
    )
    add_node_badge(doc, [
        ("UNG_DUNG.01", "Rà soát & phân tích hợp đồng", COLOR_UNG_DUNG),
    ])
    add_reason(doc, "Mô tả chức năng rà soát hợp đồng tự động rõ ràng nhất trong toàn bộ 5 phỏng vấn – đây là reference tiêu biểu nhất cho UNG_DUNG.01.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 3.2", color=COLOR_UNG_DUNG)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thứ hai là tra cứu pháp luật – chúng tôi có một chatbot AI được huấn luyện trên cơ sở dữ liệu pháp luật Việt Nam, "
        "có thể trả lời câu hỏi pháp lý nội bộ trong vài giây.",
        "EBF5EA", RGBColor(0x37, 0x86, 0x4E)
    )
    add_node_badge(doc, [
        ("UNG_DUNG.02", "Tra cứu pháp luật thông minh", COLOR_UNG_DUNG),
    ])
    add_reason(doc, "Chatbot trả lời câu hỏi pháp lý → đây là ứng dụng tra cứu thông minh điển hình.")
    add_memo(doc, "Cụm 'trong vài giây' thể hiện lợi thế tốc độ – có thể liên kết với LOI_ICH.01 khi viết báo cáo.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 3.3", color=COLOR_UNG_DUNG)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thứ ba là theo dõi sự tuân thủ – AI tự động cập nhật khi có văn bản pháp luật mới và cảnh báo nếu có nội dung "
        "liên quan đến hoạt động kinh doanh của chúng tôi.",
        "EBF5EA", RGBColor(0x37, 0x86, 0x4E)
    )
    add_node_badge(doc, [
        ("UNG_DUNG.04", "Theo dõi tuân thủ & cảnh báo", COLOR_UNG_DUNG),
    ])
    add_reason(doc, "Tự động cập nhật văn bản mới + cảnh báo → đây chính xác là compliance monitoring.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 3.4", color=COLOR_LOI_ICH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Nhờ vậy, nhóm pháp chế của tôi tiết kiệm được khoảng 40% thời gian làm việc thủ công.",
        "FFFAEB", RGBColor(0xBF, 0x85, 0x00)
    )
    add_node_badge(doc, [
        ("LOI_ICH.01", "Tiết kiệm thời gian", COLOR_LOI_ICH),
    ])
    add_reason(doc, "Con số 40% là bằng chứng định lượng về tiết kiệm thời gian – reference rất có giá trị cho LOI_ICH.01.")
    add_memo(doc, "40% tiết kiệm thời gian – ghi lại để so sánh với số liệu của NTL_02 (4 giờ → 45 phút) và NTL_05 (3-4 tuần → 7-10 ngày) khi phân tích cross-case.")

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q4
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q4 – Đánh giá độ chính xác và tin cậy của AI")
    add_body(doc, "Câu hỏi PVV: Ông đánh giá thế nào về độ chính xác và độ tin cậy của các công cụ AI này?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 4.1", color=COLOR_THACH_THUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Tôi muốn thành thật rằng AI không hoàn hảo. Trong những tháng đầu triển khai, chúng tôi gặp khá nhiều trường hợp "
        "AI hiểu sai ngữ cảnh pháp lý, đặc biệt với các điều khoản đặc thù của luật Việt Nam mà không có nhiều dữ liệu huấn luyện.",
        "FDF2EE", RGBColor(0xC0, 0x50, 0x20)
    )
    add_node_badge(doc, [
        ("THACH_THUC.03", "AI không hiểu ngữ cảnh", COLOR_THACH_THUC),
        ("THACH_THUC.01", "Hạn chế dữ liệu tiếng Việt", COLOR_THACH_THUC),
    ])
    add_reason(doc,
        "Câu 'AI hiểu sai ngữ cảnh pháp lý' → THACH_THUC.03. "
        "Cụm 'các điều khoản đặc thù của luật Việt Nam mà không có nhiều dữ liệu huấn luyện' → THACH_THUC.01 (thiếu dữ liệu TV). "
        "Hai vấn đề xuất hiện trong cùng một câu → multiple coding."
    )
    add_warning(doc, "Đây là đoạn khó: THACH_THUC.01 và THACH_THUC.03 có thể bị nhầm. THACH_THUC.01 = thiếu dữ liệu huấn luyện. THACH_THUC.03 = AI có đủ dữ liệu nhưng vẫn thiếu hiểu biết ngữ cảnh. Ở đây cả hai đều hiện diện.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 4.2", color=COLOR_LOI_ICH)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Sau khi chúng tôi bổ sung thêm dữ liệu đặc thù ngành và tinh chỉnh mô hình, độ chính xác đã cải thiện đáng kể "
        "– tôi ước tính khoảng 85-90% cho việc nhận diện rủi ro cơ bản.",
        "FFFAEB", RGBColor(0xBF, 0x85, 0x00)
    )
    add_node_badge(doc, [
        ("LOI_ICH.02", "Nâng cao chất lượng & độ chính xác", COLOR_LOI_ICH),
    ])
    add_reason(doc, "Con số 85-90% là bằng chứng định lượng về chất lượng AI sau khi tinh chỉnh → LOI_ICH.02.")
    add_memo(doc, "NTL nói 'ước tính' – không phải đo lường chính thức. Ghi chú hạn chế về độ tin cậy của con số này khi trình bày.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 4.3", color=COLOR_DAO_DUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Quan trọng là chúng tôi luôn có luật sư con người xem xét lại kết quả của AI trước khi ra quyết định. "
        "AI là công cụ hỗ trợ, không phải người ra quyết định.",
        "FDEAEA", RGBColor(0x96, 0x23, 0x1C)
    )
    add_node_badge(doc, [
        ("DAO_DUC.01", "Nguyên tắc AI hỗ trợ – con người quyết định", COLOR_DAO_DUC),
    ])
    add_reason(doc, "Câu 'AI là công cụ hỗ trợ, không phải người ra quyết định' là phát biểu nguyên tắc đạo đức cốt lõi nhất → DAO_DUC.01. Đây là trích dẫn tiêu biểu nhất cho node này.")
    add_memo(doc, "Nguyên tắc này xuất hiện ở tất cả 5 NTL. Khi phân tích, đây sẽ là theme đồng thuận mạnh nhất.")

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q5
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q5 – Những thách thức lớn nhất khi triển khai AI")
    add_body(doc, "Câu hỏi PVV: Đâu là những thách thức lớn nhất mà ông gặp phải khi triển khai AI trong công tác pháp chế?", italic=True)
    doc.add_paragraph()
    add_body(doc, "★ Câu trả lời Q5 chứa nhiều node Thách thức nhất – cần đọc kỹ từng mệnh đề.", bold=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 5.1", color=COLOR_THACH_THUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thách thức đầu tiên là về dữ liệu. Dữ liệu pháp lý bằng tiếng Việt còn rất hạn chế, "
        "và nhiều văn bản pháp luật Việt Nam chưa được số hóa đầy đủ.",
        "FDF2EE", RGBColor(0xC0, 0x50, 0x20)
    )
    add_node_badge(doc, [
        ("THACH_THUC.01", "Hạn chế dữ liệu tiếng Việt", COLOR_THACH_THUC),
    ])
    add_reason(doc, "Nêu rõ hai khía cạnh của vấn đề dữ liệu: (1) thiếu corpus pháp lý TV, (2) văn bản chưa số hóa. Cả hai đều thuộc THACH_THUC.01.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 5.2", color=COLOR_THACH_THUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thứ hai là vấn đề bảo mật – chúng tôi không thể đưa hợp đồng bảo mật lên các dịch vụ AI đám mây công cộng, "
        "nên phải đầu tư vào hạ tầng AI tại chỗ, chi phí khá cao.",
        "FDF2EE", RGBColor(0xC0, 0x50, 0x20)
    )
    add_node_badge(doc, [
        ("THACH_THUC.02", "Bảo mật & an ninh thông tin", COLOR_THACH_THUC),
        ("THACH_THUC.05", "Chi phí & hạ tầng", COLOR_THACH_THUC),
    ])
    add_reason(doc,
        "Đoạn này chứa hai vấn đề liên quan nhưng khác nhau: "
        "bảo mật (không thể dùng cloud) → THACH_THUC.02; "
        "chi phí đầu tư hạ tầng on-premise → THACH_THUC.05."
    )
    add_warning(doc, "THACH_THUC.02 và THACH_THUC.05 hay bị gộp nhầm. Bảo mật = vấn đề an toàn thông tin. Chi phí = vấn đề tài chính. Ở đây câu 'không thể đưa lên cloud' = bảo mật; câu 'phải đầu tư hạ tầng, chi phí cao' = chi phí.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 5.3", color=COLOR_THACH_THUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Thứ ba là yếu tố con người – ban đầu một số thành viên trong nhóm lo ngại AI sẽ thay thế việc làm của họ, "
        "tạo ra sự kháng cự.",
        "FDF2EE", RGBColor(0xC0, 0x50, 0x20)
    )
    add_node_badge(doc, [
        ("THACH_THUC.04", "Kháng cự & tâm lý nhân sự", COLOR_THACH_THUC),
    ])
    add_reason(doc, "Lo ngại mất việc, kháng cự thay đổi → đây là vấn đề tâm lý và văn hóa tổ chức điển hình.")
    add_memo(doc, "NTL dùng cụm 'yếu tố con người' – thú vị khi đặt cạnh nguyên tắc 'con người quyết định' ở ĐOẠN 4.3. Hai ý nghĩa của 'con người' rất khác nhau.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 5.4", color=COLOR_TUONG_LAI)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Chúng tôi phải mất nhiều thời gian để thay đổi tư duy, giúp họ hiểu rằng AI giải phóng họ khỏi công việc lặp lại "
        "để tập trung vào những vấn đề phức tạp hơn.",
        "F2F2F2", RGBColor(0x40, 0x40, 0x40)
    )
    add_node_badge(doc, [
        ("TUONG_LAI.01", "Thay đổi vai trò luật sư / pháp chế", COLOR_TUONG_LAI),
        ("THACH_THUC.04", "Kháng cự & tâm lý nhân sự", COLOR_THACH_THUC),
    ])
    add_reason(doc,
        "Câu này liên kết với THACH_THUC.04 (tiếp theo từ kháng cự) nhưng đồng thời thể hiện quan điểm về vai trò mới của pháp chế (tập trung vào vấn đề phức tạp hơn) → TUONG_LAI.01."
    )

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q6
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q6 – Khung pháp lý và khoảng trống pháp lý")
    add_body(doc, "Câu hỏi PVV: Ông có ý kiến gì về khung pháp lý hiện tại liên quan đến việc sử dụng AI trong lĩnh vực pháp lý tại Việt Nam không?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 6.1", color=COLOR_PHAP_LY)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Đây là một khoảng trống đáng lo ngại. Hiện tại Việt Nam chưa có quy định cụ thể nào điều chỉnh việc sử dụng AI "
        "trong hành nghề pháp lý hay pháp chế doanh nghiệp. Chúng tôi hoạt động trong một vùng xám pháp lý.",
        "F5EDFB", RGBColor(0x5B, 0x2C, 0x8D)
    )
    add_node_badge(doc, [
        ("PHAP_LY.01", "Khoảng trống pháp lý tại Việt Nam", COLOR_PHAP_LY),
    ])
    add_reason(doc, "Cụm 'vùng xám pháp lý' là ẩn dụ mạnh về khoảng trống pháp lý – đây là reference tiêu biểu nhất cho PHAP_LY.01.")
    add_memo(doc, "'Vùng xám pháp lý' – đây là cụm từ đáng trích dẫn trực tiếp trong báo cáo.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 6.2", color=COLOR_PHAP_LY)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Về trách nhiệm pháp lý – nếu AI đưa ra kết quả sai và gây thiệt hại, trách nhiệm thuộc về ai? "
        "Về bảo mật dữ liệu – việc xử lý hợp đồng bằng AI phải tuân thủ Nghị định 13/2023 về bảo vệ dữ liệu cá nhân như thế nào?",
        "F5EDFB", RGBColor(0x5B, 0x2C, 0x8D)
    )
    add_node_badge(doc, [
        ("PHAP_LY.02", "Trách nhiệm pháp lý của AI", COLOR_PHAP_LY),
        ("THACH_THUC.02", "Bảo mật & an ninh thông tin", COLOR_THACH_THUC),
    ])
    add_reason(doc,
        "Câu 1 đặt câu hỏi về quy trách nhiệm → PHAP_LY.02. "
        "Câu 2 nêu vấn đề tuân thủ Nghị định 13/2023 khi dùng AI → đây vừa là vấn đề pháp lý (PHAP_LY.01) vừa là thách thức bảo mật (THACH_THUC.02). "
        "Ưu tiên PHAP_LY.02 và THACH_THUC.02 vì đây là hai ý cụ thể hơn."
    )
    add_warning(doc, "Đây là đoạn phức tạp nhất trong Q6. Không cần phải chọn đúng 100% – ghi Memo nếu không chắc.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 6.3", color=COLOR_PHAP_LY)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Tôi mong muốn các cơ quan quản lý sớm ban hành hướng dẫn cụ thể để doanh nghiệp có thể yên tâm đầu tư và triển khai.",
        "F5EDFB", RGBColor(0x5B, 0x2C, 0x8D)
    )
    add_node_badge(doc, [
        ("PHAP_LY.03", "Kiến nghị chính sách", COLOR_PHAP_LY),
    ])
    add_reason(doc, "Câu này thể hiện mong muốn / kiến nghị trực tiếp với cơ quan nhà nước → PHAP_LY.03.")

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q7
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q7 – Nhìn về tương lai AI trong pháp chế")
    add_body(doc, "Câu hỏi PVV: Nhìn về tương lai, ông thấy AI sẽ thay đổi công tác pháp chế doanh nghiệp như thế nào?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 7.1", color=COLOR_TUONG_LAI)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Tôi tin rằng trong 5-10 năm tới, AI sẽ đảm nhiệm phần lớn công việc pháp chế mang tính lặp lại và tiêu chuẩn hóa "
        "– soạn thảo hợp đồng theo mẫu, rà soát tuân thủ, tra cứu pháp luật.",
        "F2F2F2", RGBColor(0x40, 0x40, 0x40)
    )
    add_node_badge(doc, [
        ("TUONG_LAI.01", "Thay đổi vai trò luật sư / pháp chế", COLOR_TUONG_LAI),
    ])
    add_reason(doc, "Dự đoán AI đảm nhiệm công việc lặp lại → thể hiện sự dịch chuyển vai trò. Cụm '5-10 năm tới' = tầm nhìn tương lai rõ ràng.")
    add_memo(doc, "NTL liệt kê: soạn thảo theo mẫu, rà soát tuân thủ, tra cứu pháp luật = ba ứng dụng AI đã được nhắc ở Q3. Sự nhất quán này đáng ghi chú.")

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 7.2", color=COLOR_TUONG_LAI)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Vai trò của luật sư và chuyên gia pháp chế sẽ dịch chuyển sang tư vấn chiến lược, đàm phán phức tạp, "
        "và những vấn đề đòi hỏi phán đoán đạo đức và kinh nghiệm thực tiễn.",
        "F2F2F2", RGBColor(0x40, 0x40, 0x40)
    )
    add_node_badge(doc, [
        ("TUONG_LAI.01", "Thay đổi vai trò luật sư / pháp chế", COLOR_TUONG_LAI),
        ("TUONG_LAI.02", "Kỹ năng cần thiết trong thời đại AI", COLOR_TUONG_LAI),
    ])
    add_reason(doc,
        "Câu đầu nói về hướng dịch chuyển vai trò → TUONG_LAI.01. "
        "Cụm 'phán đoán đạo đức và kinh nghiệm thực tiễn' ngụ ý kỹ năng con người không thể thay thế → TUONG_LAI.02."
    )

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 7.3", color=COLOR_TUONG_LAI)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Tôi cũng kỳ vọng sẽ có các nền tảng AI pháp lý chuyên biệt cho thị trường Việt Nam, được xây dựng trên nền tảng "
        "pháp luật Việt Nam và được các cơ quan chức năng chứng nhận.",
        "F2F2F2", RGBColor(0x40, 0x40, 0x40)
    )
    add_node_badge(doc, [
        ("TUONG_LAI.05", "So sánh Việt Nam với quốc tế", COLOR_TUONG_LAI),
        ("PHAP_LY.03", "Kiến nghị chính sách", COLOR_PHAP_LY),
    ])
    add_reason(doc,
        "Kỳ vọng nền tảng AI đặc thù cho Việt Nam → ngụ ý nhận thức về khoảng cách với quốc tế → TUONG_LAI.05. "
        "Yếu tố 'được các cơ quan chức năng chứng nhận' = mong muốn chính sách nhà nước → PHAP_LY.03."
    )

    add_divider(doc)

    # ══════════════════════════════════════════════════════════════════════
    # Q8
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "Q8 – Lời nhắn nhủ cuối")
    add_body(doc, "Câu hỏi PVV: Ông có muốn bổ sung điều gì không?", italic=True)
    doc.add_paragraph()

    add_h3(doc, "ĐOẠN 8.1", color=COLOR_DAO_DUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Doanh nghiệp cần đầu tư nghiêm túc vào việc lựa chọn công cụ phù hợp, đào tạo nhân sự, "
        "và thiết lập quy trình kiểm soát chất lượng.",
        "FDEAEA", RGBColor(0x96, 0x23, 0x1C)
    )
    add_node_badge(doc, [
        ("TUONG_LAI.02", "Kỹ năng cần thiết trong thời đại AI", COLOR_TUONG_LAI),
        ("TUONG_LAI.03", "Đào tạo & cải cách giáo dục luật", COLOR_TUONG_LAI),
    ])
    add_reason(doc,
        "Lời khuyên thực tiễn: đào tạo nhân sự → TUONG_LAI.03; "
        "quy trình kiểm soát chất lượng + lựa chọn công cụ = kỹ năng cần có → TUONG_LAI.02."
    )

    doc.add_paragraph()
    add_h3(doc, "ĐOẠN 8.2", color=COLOR_DAO_DUC)
    add_data_box(doc,
        "Đoạn cần bôi chọn trong NVivo:",
        "Quan trọng hơn là không bao giờ để AI ra quyết định pháp lý mà không có sự giám sát của con người. "
        "Pháp lý là lĩnh vực có hậu quả nghiêm trọng nếu sai sót, nên ngưỡng kiểm soát phải rất cao.",
        "FDEAEA", RGBColor(0x96, 0x23, 0x1C)
    )
    add_node_badge(doc, [
        ("DAO_DUC.01", "Nguyên tắc AI hỗ trợ – con người quyết định", COLOR_DAO_DUC),
    ])
    add_reason(doc, "Phát biểu nguyên tắc đạo đức một lần nữa ở phần kết – NTL nhấn mạnh lần thứ hai trong cùng bài phỏng vấn. Điều này cho thấy mức độ quan trọng của nguyên tắc này với NTL_01.")
    add_memo(doc, "NTL nhắc DAO_DUC.01 hai lần (ĐOẠN 4.3 và ĐOẠN 8.2). Sự lặp lại chủ động này = niềm tin sâu sắc, không phải câu trả lời xã giao. Đáng nhấn mạnh trong báo cáo.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # TỔNG KẾT
    # ══════════════════════════════════════════════════════════════════════
    add_h2(doc, "TỔNG KẾT – BẢNG MÃ HÓA PV01")

    add_body(doc, "Sau khi hoàn thành code selection PV01, kiểm tra lại theo bảng sau:")
    doc.add_paragraph()

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    sh = summary_table.rows[0].cells
    for i, h in enumerate(["Node", "Số đoạn", "Đoạn số", "Ghi chú"]):
        sh[i].text = h
        sh[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(sh[i], "2E74B5")
        sh[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    summary_data = [
        ("BCLC.01 – Áp lực công việc",        "3", "1.2, 2.2, 5.4*",   "Đoạn 2.2 là reference mạnh nhất",         "EBF3FC"),
        ("BCLC.03 – Lộ trình triển khai",      "2", "2.1, 2.3",         "Mốc thời gian cụ thể",                    "EBF3FC"),
        ("BCLC.04 – Đặc thù ngành",            "2", "1.1, 1.2",         "Fintech thay đổi nhanh",                  "EBF3FC"),
        ("UNG_DUNG.01 – Rà soát hợp đồng",     "2", "2.3*, 3.1",        "Đoạn 3.1 là reference tiêu biểu nhất",   "EBF5EA"),
        ("UNG_DUNG.02 – Tra cứu pháp luật",    "1", "3.2",              "Chatbot pháp lý",                         "EBF5EA"),
        ("UNG_DUNG.04 – Theo dõi tuân thủ",    "1", "3.3",              "Compliance monitoring",                   "EBF5EA"),
        ("LOI_ICH.01 – Tiết kiệm thời gian",   "1", "3.4",              "Con số 40% – định lượng",                 "FFFAEB"),
        ("LOI_ICH.02 – Chất lượng/Chính xác",  "1", "4.2",              "Con số 85-90% – định lượng",              "FFFAEB"),
        ("THACH_THUC.01 – Dữ liệu TV",         "2", "4.1*, 5.1",        "Đoạn 5.1 rõ ràng hơn",                   "FDF2EE"),
        ("THACH_THUC.02 – Bảo mật",            "2", "5.2*, 6.2*",       "Cloud vs on-premise",                     "FDF2EE"),
        ("THACH_THUC.03 – Ngữ cảnh",           "1", "4.1*",             "AI hiểu sai ngữ cảnh",                    "FDF2EE"),
        ("THACH_THUC.04 – Kháng cự nhân sự",   "2", "5.3, 5.4*",        "Lo ngại mất việc",                        "FDF2EE"),
        ("THACH_THUC.05 – Chi phí/Hạ tầng",    "1", "5.2*",             "Đầu tư on-premise",                       "FDF2EE"),
        ("PHAP_LY.01 – Khoảng trống pháp lý",  "1", "6.1",              "'Vùng xám pháp lý' – trích dẫn mạnh",    "F5EDFB"),
        ("PHAP_LY.02 – Trách nhiệm AI",        "1", "6.2*",             "Ai chịu trách nhiệm khi AI sai?",         "F5EDFB"),
        ("PHAP_LY.03 – Kiến nghị chính sách",  "2", "6.3, 7.3*",        "Mong muốn hướng dẫn từ nhà nước",        "F5EDFB"),
        ("DAO_DUC.01 – Human in the loop",      "2", "4.3, 8.2",         "NTL nhắc 2 lần – niềm tin sâu sắc",      "FDEAEA"),
        ("TUONG_LAI.01 – Thay đổi vai trò",    "3", "5.4*, 7.1, 7.2*",  "Dịch chuyển sang tư vấn chiến lược",     "F2F2F2"),
        ("TUONG_LAI.02 – Kỹ năng AI",          "2", "7.2*, 8.1*",       "Phán đoán đạo đức + quy trình KS",       "F2F2F2"),
        ("TUONG_LAI.03 – Đào tạo",             "1", "8.1*",             "Đào tạo nhân sự",                         "F2F2F2"),
        ("TUONG_LAI.05 – So sánh quốc tế",     "1", "7.3*",             "AI chuyên biệt cho VN",                   "F2F2F2"),
    ]

    for (node, so_doan, doan_so, ghi_chu, bg) in summary_data:
        row = summary_table.add_row().cells
        row[0].text = node
        row[1].text = so_doan
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].text = doan_so
        row[3].text = ghi_chu
        for cell in row:
            set_cell_bg(cell, bg)

    doc.add_paragraph()
    add_body(doc, "* = đoạn có multiple coding (được mã hóa vào nhiều node).", italic=True)
    doc.add_paragraph()

    # Thống kê
    add_h3(doc, "Thống kê nhanh", color=(0x1F, 0x49, 0x7D))
    stats = [
        ("Tổng số đoạn mã hóa đề xuất", "21 đoạn"),
        ("Tổng số lần gán node (references)", "27 lần (do multiple coding)"),
        ("Số node được kích hoạt", "21 / 35 nodes (60%)"),
        ("Node xuất hiện nhiều nhất", "TUONG_LAI.01, THACH_THUC nhóm (5 nodes)"),
        ("Node không xuất hiện trong PV01", "UNG_DUNG.05-08, LOI_ICH.03-04, PHAP_LY.04, DAO_DUC.02-04"),
    ]
    for k, v in stats:
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)

    doc.save(os.path.join(OUTPUT_DIR, "HUONG_DAN_CodeSelection_PV01.docx"))
    print("Đã tạo: HUONG_DAN_CodeSelection_PV01.docx")


if __name__ == "__main__":
    create_pv01_coding_guide()
