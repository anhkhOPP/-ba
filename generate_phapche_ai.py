# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), val.get('sz', '4'))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_header_cell(cell, text, font_size=9, bold=True, bg_color="1F4E79", font_color="FFFFFF"):
    set_cell_bg(cell, bg_color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(font_color)
    run.font.name = 'Times New Roman'

def add_body_cell(cell, text, font_size=8.5, bold=False, bg_color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg_color:
        set_cell_bg(cell, bg_color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'

def set_col_widths(table, widths):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = widths[j]

doc = Document()

# ===== PAGE SETUP =====
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ===== TITLE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ỨNG DỤNG TRÍ TUỆ NHÂN TẠO (AI) TRONG CÔNG VIỆC\nPHÁP CHẾ NGÂN HÀNG TECHCOMBANK")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Tổng hợp: Bộ phận Pháp chế – Techcombank | Tháng 4/2026")
run2.font.size = Pt(10)
run2.italic = True
run2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
run2.font.name = 'Times New Roman'

doc.add_paragraph()

# ===== DATA =====
# Structure: (nhom, stt, loai_cv, mo_ta, quy_trinh, ai_cong_nghe)

data = [
    # NHOM 1
    ("NHÓM 1: SOẠN THẢO & RÀ SOÁT HỢP ĐỒNG / TÀI LIỆU PHÁP LÝ", "1.1",
     "Soạn thảo hợp đồng",
     "Soạn thảo hợp đồng tín dụng, hợp đồng bảo lãnh, hợp đồng thế chấp, hợp đồng dịch vụ ngân hàng, hợp đồng cung cấp phần mềm, NDA, hợp đồng lao động... từ đầu hoặc từ mẫu có sẵn.",
     "1. Thu thập yêu cầu từ bộ phận nghiệp vụ\n2. Xác định loại hợp đồng, khung pháp lý áp dụng\n3. Tra cứu mẫu/biểu mẫu nội bộ\n4. Soạn thảo nội dung\n5. Rà soát nội bộ\n6. Đàm phán, chỉnh sửa với đối tác\n7. Phê duyệt, ký kết\n8. Lưu trữ hệ thống",
     "• Generative AI (GPT-4o, Claude 3.5, Gemini 1.5): sinh bản thảo hợp đồng tự động từ template + yêu cầu đầu vào\n• Contract Lifecycle Management AI: DocuSign CLM, Ironclad, Juro\n• Harvey AI, CoCounsel (AI pháp lý chuyên biệt)\n• Microsoft Copilot for Word: hỗ trợ soạn thảo trực tiếp trong Word"),

    ("", "1.2",
     "Rà soát hợp đồng do đối tác soạn",
     "Đọc, phân tích hợp đồng do đối tác/khách hàng/nhà cung cấp cung cấp; xác định điều khoản bất lợi, rủi ro pháp lý, điều khoản bất thường, thiếu sót so với chuẩn nội bộ.",
     "1. Nhận hợp đồng từ đối tác\n2. Xác định tiêu chí rà soát (phạt, bồi thường, chấm dứt, giải quyết tranh chấp, luật áp dụng...)\n3. Đối chiếu với chính sách nội bộ và quy định pháp luật\n4. Lập danh sách vấn đề & đề xuất redline\n5. Phản hồi/đàm phán với đối tác\n6. Lưu trữ phiên bản đã thống nhất",
     "• AI Contract Review: Luminance, Kira Systems, Lexion, ThoughtRiver\n• Harvey AI, CoCounsel: phân tích điều khoản rủi ro\n• Redlining tự động: Microsoft Copilot for Word\n• LegalSifter: gợi ý điều khoản thiếu và rủi ro theo checklist"),

    ("", "1.3",
     "Chuẩn hóa mẫu hợp đồng / thư viện điều khoản",
     "Cập nhật, chuẩn hóa thư viện mẫu hợp đồng và clause library theo quy định pháp luật mới hoặc chính sách nội bộ, đảm bảo nhất quán trong toàn hệ thống.",
     "1. Rà soát toàn bộ mẫu hợp đồng hiện hành\n2. Đối chiếu với thay đổi pháp luật gần nhất\n3. Xác định điều khoản cần cập nhật/bổ sung/loại bỏ\n4. Soạn phiên bản mới\n5. Phê duyệt cấp có thẩm quyền\n6. Cập nhật vào hệ thống quản lý tài liệu",
     "• AI Clause Library: Ironclad, Juro, Contractbook\n• Semantic Search AI: tìm kiếm điều khoản tương đồng trong kho hợp đồng\n• NLP clustering: phân loại và nhóm điều khoản tự động"),

    ("", "1.4",
     "Dịch thuật hợp đồng / tài liệu pháp lý",
     "Dịch hợp đồng, điều lệ, chính sách, báo cáo pháp lý từ tiếng Anh sang tiếng Việt hoặc ngược lại, đảm bảo thuật ngữ pháp lý chính xác và nhất quán.",
     "1. Nhận tài liệu cần dịch\n2. Xác định thuật ngữ chuyên ngành đặc thù\n3. Dịch sơ bộ bằng công cụ AI\n4. Hiệu đính thuật ngữ pháp lý ngân hàng\n5. Kiểm tra tính nhất quán toàn văn bản\n6. Phê duyệt và lưu trữ",
     "• DeepL Pro (dịch pháp lý chất lượng cao)\n• ChatGPT / Claude: dịch + giải thích thuật ngữ\n• Glossary AI: xây dựng từ điển thuật ngữ nội bộ\n• SDL Trados tích hợp AI (CAT tool)"),

    # NHOM 2
    ("NHÓM 2: NGHIÊN CỨU & TRA CỨU PHÁP LUẬT", "2.1",
     "Tra cứu văn bản pháp luật, quy định",
     "Tra cứu, tổng hợp các quy định pháp luật liên quan đến hoạt động ngân hàng: Luật Các TCTD, Luật Chứng khoán, Luật PCRT, Thông tư NHNN, Nghị định Chính phủ...",
     "1. Nhận yêu cầu tra cứu từ bộ phận nghiệp vụ/ban lãnh đạo\n2. Xác định vấn đề pháp lý cần tra cứu\n3. Xác định nguồn văn bản (Công báo, VBQPPL, cơ sở dữ liệu pháp luật)\n4. Tra cứu, sàng lọc văn bản liên quan\n5. Tổng hợp, phân tích nội dung\n6. Lập ý kiến pháp lý/memo\n7. Gửi phản hồi cho bộ phận yêu cầu",
     "• AI Legal Research: Westlaw Edge AI, LexisNexis+AI, Fastcase AI\n• Chatbot pháp lý nội bộ tích hợp RAG (Retrieval-Augmented Generation)\n• LLM kết hợp vector database (FAISS, Pinecone): tra cứu toàn văn văn bản pháp luật Việt Nam\n• LegalMind Vietnam (nếu triển khai): AI tra cứu pháp luật Việt Nam"),

    ("", "2.2",
     "Theo dõi, cập nhật thay đổi pháp luật",
     "Theo dõi thường xuyên các thay đổi về quy định pháp luật, Thông tư NHNN, Nghị định, Quyết định có ảnh hưởng đến hoạt động ngân hàng; cảnh báo kịp thời cho các bộ phận liên quan.",
     "1. Thiết lập danh mục văn bản cần theo dõi\n2. Giám sát Công báo, website NHNN, Chính phủ, Quốc hội\n3. Phân loại mức độ ảnh hưởng (high/medium/low)\n4. Tóm tắt nội dung thay đổi\n5. Gửi cảnh báo/bản tin nội bộ\n6. Đề xuất điều chỉnh quy trình, chính sách nội bộ",
     "• AI Regulatory Monitoring: Compliance.ai, Ascent RegTech, Cube RegTech\n• Web scraping + NLP: tự động thu thập và phân loại văn bản mới\n• Automated alert system: gửi email/notification khi có văn bản mới liên quan\n• LLM tóm tắt: tự động tóm tắt nội dung thay đổi chính"),

    ("", "2.3",
     "Nghiên cứu án lệ, tiền lệ pháp",
     "Tra cứu, phân tích án lệ của TAND Tối cao, phán quyết trọng tài, tiền lệ liên quan đến tranh chấp ngân hàng – tín dụng – bảo đảm để hỗ trợ xử lý vụ việc.",
     "1. Xác định vấn đề pháp lý cần tra cứu án lệ\n2. Tra cứu trên cơ sở dữ liệu án lệ\n3. Phân tích, tóm tắt án lệ có liên quan\n4. Đánh giá khả năng áp dụng vào vụ việc cụ thể\n5. Lập memo phân tích\n6. Lưu trữ vào thư viện án lệ nội bộ",
     "• AI Case Law Analysis: Casetext CARA AI, Westlaw AI\n• LLM + RAG: xây dựng chatbot tra cứu án lệ nội bộ\n• Semantic similarity search: tìm án lệ tương tự theo tình tiết vụ việc"),

    ("", "2.4",
     "Soạn thảo ý kiến pháp lý (Legal Opinion / Legal Memo)",
     "Soạn thảo ý kiến pháp lý trả lời câu hỏi của các bộ phận kinh doanh, ban lãnh đạo về tính hợp pháp của sản phẩm, giao dịch, quy trình, cơ cấu đầu tư...",
     "1. Nhận câu hỏi/yêu cầu từ bộ phận\n2. Xác định vấn đề pháp lý cốt lõi\n3. Nghiên cứu quy định áp dụng\n4. Phân tích rủi ro pháp lý\n5. Soạn thảo ý kiến pháp lý\n6. Rà soát nội bộ và ký duyệt\n7. Gửi phản hồi",
     "• LLM (GPT-4o, Claude): hỗ trợ soạn thảo cấu trúc memo, gợi ý luận điểm\n• Harvey AI: soạn thảo legal memo chuyên biệt\n• AI outline generator: tạo khung phân tích pháp lý\n• Grammarly / LanguageTool: kiểm tra văn phong, lỗi chính tả"),

    # NHOM 3
    ("NHÓM 3: TUÂN THỦ PHÁP LÝ (COMPLIANCE)", "3.1",
     "Rà soát tuân thủ nội bộ (Internal Compliance Review)",
     "Rà soát định kỳ quy trình, chính sách, sản phẩm nội bộ để đảm bảo tuân thủ quy định pháp luật hiện hành và quy định nội bộ của NHNN, Techcombank.",
     "1. Lập kế hoạch rà soát định kỳ\n2. Thu thập tài liệu quy trình/chính sách cần rà soát\n3. Đối chiếu với quy định pháp luật áp dụng\n4. Xác định khoảng cách tuân thủ (compliance gap)\n5. Lập báo cáo kết quả rà soát\n6. Đề xuất biện pháp khắc phục\n7. Theo dõi thực hiện khắc phục",
     "• AI Compliance Mapping: Clausematch, Ascent, 6clicks\n• LLM + Rule-based AI: đối chiếu tự động quy trình nội bộ với văn bản pháp luật\n• GRC Platform tích hợp AI: ServiceNow GRC, MetricStream\n• Automated gap analysis: phát hiện điểm không tuân thủ"),

    ("", "3.2",
     "Xây dựng, cập nhật quy định, quy trình nội bộ",
     "Soạn thảo, rà soát, cập nhật các quy chế, quy trình, hướng dẫn nội bộ (credit policy, KYC policy, data privacy policy, anti-corruption policy...) theo yêu cầu pháp luật mới.",
     "1. Xác định văn bản nội bộ cần xây dựng/cập nhật\n2. Thu thập yêu cầu từ các bộ phận liên quan\n3. Nghiên cứu quy định pháp luật, thực tiễn ngành\n4. Soạn thảo\n5. Lấy ý kiến các bộ phận\n6. Tổng hợp, chỉnh sửa\n7. Phê duyệt ban hành\n8. Truyền thông nội bộ",
     "• Generative AI: soạn thảo nhánh ban đầu từ yêu cầu\n• Document comparison AI: so sánh phiên bản cũ và mới\n• Policy management platform tích hợp AI: PowerDMS, Ncontracts\n• Chatbot AI: trả lời thắc mắc nhân viên về chính sách"),

    ("", "3.3",
     "Phòng chống rửa tiền (AML/CFT Compliance)",
     "Rà soát, cập nhật quy trình AML/CFT; hỗ trợ xử lý các trường hợp cảnh báo giao dịch đáng ngờ (STR); đảm bảo tuân thủ Luật PCRT và quy định NHNN về AML.",
     "1. Giám sát cảnh báo giao dịch đáng ngờ từ hệ thống\n2. Phân tích thông tin khách hàng, giao dịch\n3. Tra cứu danh sách cấm vận (sanctions screening)\n4. Xác định có/không có dấu hiệu rửa tiền\n5. Lập báo cáo STR gửi Cục PCRT\n6. Lưu hồ sơ\n7. Cập nhật quy trình định kỳ",
     "• AI Transaction Monitoring: NICE Actimize, Oracle FCCM, Temenos Financial Crime\n• Machine Learning anomaly detection: phát hiện giao dịch bất thường\n• NLP entity extraction: trích xuất thông tin thực thể từ STR\n• Sanctions screening AI: World-Check One, Dow Jones Risk"),

    ("", "3.4",
     "Tuân thủ bảo vệ dữ liệu cá nhân (PDPA/Data Privacy)",
     "Đảm bảo hoạt động xử lý dữ liệu cá nhân của ngân hàng tuân thủ Nghị định 13/2023/NĐ-CP, PDPA; soạn thảo điều khoản bảo mật, đánh giá tác động quyền riêng tư (DPIA).",
     "1. Rà soát toàn bộ luồng dữ liệu cá nhân trong ngân hàng\n2. Xác định căn cứ pháp lý xử lý dữ liệu\n3. Thực hiện DPIA cho hoạt động rủi ro cao\n4. Soạn thảo/cập nhật Privacy Notice, Consent Form\n5. Xây dựng quy trình xử lý yêu cầu của chủ thể dữ liệu\n6. Báo cáo sự cố dữ liệu (nếu có)",
     "• AI Data Mapping: OneTrust AI, TrustArc, Securiti.ai\n• Automated DPIA tools: tự động đánh giá rủi ro quyền riêng tư\n• LLM: soạn thảo Privacy Notice, Consent Form\n• Data discovery AI: tự động phát hiện dữ liệu cá nhân trong hệ thống"),

    # NHOM 4
    ("NHÓM 4: QUẢN LÝ TRANH TỤNG & XỬ LÝ NỢ PHÁP LÝ", "4.1",
     "Xử lý tài sản bảo đảm (thu hồi nợ)",
     "Thực hiện các thủ tục pháp lý để xử lý tài sản bảo đảm (TSBĐ) khi khách hàng vi phạm nghĩa vụ: phát mại, chuyển nhượng, nhận chính TSBĐ để thay thế nghĩa vụ.",
     "1. Xác định điều kiện xử lý TSBĐ\n2. Gửi thông báo cho bên bảo đảm\n3. Thẩm định lại giá trị TSBĐ\n4. Lựa chọn phương thức xử lý\n5. Thực hiện thủ tục pháp lý (công chứng, đấu giá, sang tên...)\n6. Thu tiền, tất toán khoản nợ\n7. Giải chấp, lưu hồ sơ",
     "• AI Document Automation: tự động tạo hồ sơ yêu cầu xử lý TSBĐ\n• AI Case Prioritization: phân loại mức độ ưu tiên xử lý theo giá trị/rủi ro\n• Property valuation AI: ước tính giá trị TSBĐ bất động sản\n• LLM: soạn thảo thông báo, văn bản yêu cầu bàn giao TSBĐ"),

    ("", "4.2",
     "Khởi kiện / Tham gia tố tụng",
     "Chuẩn bị hồ sơ khởi kiện, tham gia phiên tòa, làm việc với luật sư thuê ngoài trong các vụ tranh chấp tín dụng, tranh chấp hợp đồng, xử lý tài sản.",
     "1. Đánh giá khả năng và chiến lược kiện tụng\n2. Thu thập, sắp xếp chứng cứ\n3. Soạn đơn khởi kiện, tài liệu tố tụng\n4. Nộp hồ sơ tòa án\n5. Theo dõi tiến trình vụ kiện\n6. Tham gia phiên tòa/hòa giải\n7. Thi hành án",
     "• AI Litigation Analytics: Lex Machina, Premonition AI: dự báo kết quả tố tụng\n• AI Document Review (eDiscovery): Relativity, Logikcull, Everlaw\n• Timeline automation: tự động tạo timeline vụ kiện\n• LLM: soạn thảo đơn khởi kiện, bản bào chữa\n• AI evidence management: tổ chức, phân loại tài liệu chứng cứ"),

    ("", "4.3",
     "Hòa giải, đàm phán giải quyết tranh chấp ngoài tòa",
     "Đàm phán, thương lượng với khách hàng/đối tác để giải quyết tranh chấp thông qua hòa giải, tái cơ cấu khoản nợ, ký kết thỏa thuận hòa giải.",
     "1. Phân tích vị thế đàm phán của ngân hàng\n2. Xác định phương án đề xuất\n3. Tiến hành các phiên làm việc với đối phương\n4. Soạn thảo thỏa thuận hòa giải\n5. Ký kết, công chứng (nếu cần)\n6. Theo dõi thực hiện thỏa thuận",
     "• AI negotiation simulation: mô phỏng kịch bản đàm phán\n• Sentiment analysis AI: phân tích phản ứng đối phương\n• LLM: soạn thảo thỏa thuận hòa giải\n• AI risk-reward analysis: đánh giá phương án hòa giải vs kiện tụng"),

    ("", "4.4",
     "Theo dõi tiến độ vụ kiện, thi hành án",
     "Quản lý danh sách vụ kiện đang xử lý, theo dõi deadline tố tụng, cập nhật tiến độ thi hành án, báo cáo định kỳ cho ban lãnh đạo.",
     "1. Cập nhật trạng thái từng vụ kiện vào hệ thống\n2. Thiết lập cảnh báo deadline tố tụng\n3. Thu thập kết quả phiên tòa, quyết định thi hành án\n4. Phối hợp với cơ quan thi hành án\n5. Báo cáo tiến độ định kỳ\n6. Đóng hồ sơ khi hoàn tất",
     "• Matter Management AI: Clio, MyCase, Legal Tracker (Thomson Reuters)\n• Automated deadline tracking: cảnh báo hạn nộp hồ sơ tố tụng\n• Dashboard BI: Power BI/Tableau tích hợp data pháp lý\n• AI reporting: tự động tổng hợp báo cáo trạng thái vụ kiện"),

    # NHOM 5
    ("NHÓM 5: THẨM ĐỊNH PHÁP LÝ (DUE DILIGENCE)", "5.1",
     "Thẩm định pháp lý dự án/khách hàng doanh nghiệp",
     "Thẩm định pháp lý hồ sơ pháp nhân, tư cách chủ thể, lịch sử pháp lý, tranh chấp của khách hàng doanh nghiệp vay vốn, đối tác chiến lược, bên tham gia giao dịch M&A.",
     "1. Nhận yêu cầu thẩm định từ bộ phận tín dụng/đầu tư\n2. Thu thập hồ sơ pháp lý khách hàng\n3. Tra cứu thông tin đăng ký doanh nghiệp, lịch sử tranh chấp\n4. Phân tích cơ cấu sở hữu, người có liên quan\n5. Kiểm tra tài sản bảo đảm\n6. Lập báo cáo thẩm định pháp lý\n7. Trình cấp phê duyệt",
     "• AI Document Extraction: Kira Systems, Luminance: trích xuất thông tin từ hồ sơ pháp lý\n• Entity resolution AI: xác định mối quan hệ giữa các pháp nhân, UBO\n• AI KYC/KYB: Onfido, Jumio, Moody's KYC360\n• OCR + NLP: số hóa và phân tích hồ sơ giấy tờ\n• Adverse media screening AI: kiểm tra thông tin tiêu cực"),

    ("", "5.2",
     "Thẩm định pháp lý tài sản bảo đảm",
     "Kiểm tra tính hợp lệ pháp lý của TSBĐ (bất động sản, động sản, cổ phần...): quyền sở hữu, tình trạng pháp lý, giao dịch bảo đảm đã đăng ký, tranh chấp, phong tỏa...",
     "1. Nhận hồ sơ TSBĐ\n2. Kiểm tra giấy tờ sở hữu, pháp lý tài sản\n3. Tra cứu trung tâm đăng ký giao dịch bảo đảm\n4. Tra cứu tình trạng quy hoạch, hạn chế sử dụng\n5. Xác định rủi ro pháp lý\n6. Lập ý kiến pháp lý về TSBĐ\n7. Lưu hồ sơ thẩm định",
     "• OCR AI: số hóa giấy chứng nhận, sổ đỏ\n• AI document verification: xác minh tính xác thực của giấy tờ\n• GIS + AI: tra cứu quy hoạch, tình trạng pháp lý bất động sản\n• Automated registry search: tự động tra cứu đăng ký bảo đảm"),

    ("", "5.3",
     "Thẩm định pháp lý giao dịch M&A / đầu tư",
     "Thực hiện legal due diligence trong giao dịch mua bán, sáp nhập doanh nghiệp, đầu tư vốn; rà soát hợp đồng, tranh chấp, nghĩa vụ tiềm tàng của mục tiêu M&A.",
     "1. Thiết lập phạm vi thẩm định (scope of DD)\n2. Chuẩn bị checklist thẩm định\n3. Thu thập tài liệu từ data room\n4. Phân tích, rà soát hồ sơ pháp lý\n5. Xác định rủi ro và nghĩa vụ tiềm tàng\n6. Soạn thảo báo cáo DD\n7. Hỗ trợ đàm phán SPA/SHA",
     "• Virtual Data Room AI: Intralinks AI, Datasite Acquire\n• AI DD automation: Kira Systems, Luminance: đọc và phân loại hàng nghìn tài liệu\n• Risk flagging AI: tự động đánh dấu điểm rủi ro\n• LLM: soạn thảo báo cáo DD tóm tắt\n• AI Q&A over documents: chatbot trả lời câu hỏi về data room"),

    # NHOM 6
    ("NHÓM 6: QUẢN TRỊ CÔNG TY & PHÁP LÝ NỘI BỘ", "6.1",
     "Quản lý hồ sơ pháp lý doanh nghiệp (Corporate Secretarial)",
     "Duy trì hồ sơ pháp lý của ngân hàng: Giấy phép, Điều lệ, Nghị quyết HĐQT/ĐHĐCĐ, thông báo thay đổi đăng ký kinh doanh, hồ sơ người có liên quan...",
     "1. Lưu trữ và cập nhật hồ sơ pháp lý doanh nghiệp\n2. Chuẩn bị tài liệu họp HĐQT/ĐHĐCĐ\n3. Ghi chép biên bản, nghị quyết\n4. Thực hiện thủ tục đăng ký thay đổi với cơ quan nhà nước\n5. Theo dõi hạn hạn giấy phép, nghĩa vụ báo cáo định kỳ\n6. Lưu trữ hệ thống",
     "• Board Portal AI: Diligent Boards, BoardEffect: quản lý họp và tài liệu HĐQT\n• Document management AI: iManage, NetDocuments\n• Automated deadline & obligation tracking\n• AI minutes generation: tự động tóm tắt/tạo biên bản họp từ ghi âm\n• OCR + DMS: số hóa hồ sơ giấy tờ lịch sử"),

    ("", "6.2",
     "Tư vấn pháp lý nội bộ (Internal Legal Advisory)",
     "Trả lời câu hỏi pháp lý từ các bộ phận nghiệp vụ, hỗ trợ thiết kế sản phẩm, quy trình mới đảm bảo tuân thủ; phát biểu ý kiến pháp lý trong các cuộc họp nội bộ.",
     "1. Nhận yêu cầu tư vấn\n2. Phân loại mức độ ưu tiên\n3. Nghiên cứu quy định áp dụng\n4. Chuẩn bị ý kiến pháp lý\n5. Phản hồi cho bộ phận yêu cầu\n6. Lưu hồ sơ tư vấn",
     "• Internal Legal Chatbot (RAG-based): chatbot AI tra cứu quy định nội bộ, pháp luật\n• LLM assistant: hỗ trợ soạn thảo ý kiến nhanh\n• Knowledge management AI: lưu trữ, tìm kiếm kho ý kiến pháp lý cũ\n• Ticket management + AI triage: phân loại và định tuyến yêu cầu tư vấn"),

    ("", "6.3",
     "Đào tạo pháp lý nội bộ (Legal Training)",
     "Thiết kế và triển khai chương trình đào tạo pháp lý cho cán bộ ngân hàng về quy định mới, tuân thủ AML, bảo vệ dữ liệu, phòng chống tham nhũng...",
     "1. Xác định nhu cầu đào tạo\n2. Xây dựng nội dung/curriculum\n3. Soạn thảo tài liệu học tập\n4. Tổ chức buổi đào tạo (offline/online)\n5. Kiểm tra đánh giá sau đào tạo\n6. Cập nhật tài liệu theo thay đổi pháp luật",
     "• AI content generation: tự động tạo tài liệu đào tạo, bài kiểm tra\n• LMS tích hợp AI: Coursera for Business, TalentLMS\n• AI quiz generator: tạo câu hỏi kiểm tra từ tài liệu\n• Chatbot AI: trợ lý học tập 24/7 cho nhân viên\n• Video AI (Synthesia): tạo video đào tạo pháp lý tự động"),

    # NHOM 7
    ("NHÓM 7: QUẢN LÝ HỒ SƠ & BÁO CÁO PHÁP LÝ", "7.1",
     "Quản lý, lưu trữ hồ sơ pháp lý",
     "Tổ chức, phân loại, số hóa và lưu trữ toàn bộ hồ sơ pháp lý: hợp đồng, hồ sơ tín dụng, hồ sơ tố tụng, giấy phép, ý kiến pháp lý, văn bản nội bộ...",
     "1. Thu thập hồ sơ từ các bộ phận\n2. Phân loại theo danh mục\n3. Số hóa hồ sơ giấy (scan/OCR)\n4. Lập metadata, index tìm kiếm\n5. Lưu trữ trên hệ thống DMS\n6. Thiết lập quy tắc bảo mật, phân quyền truy cập\n7. Thiết lập lịch lưu trữ và hủy tài liệu",
     "• Document Management System (DMS) + AI: iManage, NetDocuments, SharePoint Copilot\n• OCR + AI classification: tự động phân loại tài liệu theo nội dung\n• Smart tagging AI: gắn tag metadata tự động\n• AI search (semantic search): tìm kiếm tài liệu theo ngữ nghĩa\n• Retention policy automation: tự động quản lý vòng đời tài liệu"),

    ("", "7.2",
     "Báo cáo pháp lý định kỳ cho NHNN và cơ quan nhà nước",
     "Lập và nộp các báo cáo pháp lý theo quy định: báo cáo tuân thủ, báo cáo giao dịch đáng ngờ (STR), báo cáo sở hữu cổ phần, báo cáo hợp đồng lớn...",
     "1. Xác định danh mục báo cáo và deadline\n2. Thu thập dữ liệu từ các hệ thống nội bộ\n3. Xử lý, tổng hợp dữ liệu\n4. Soạn thảo báo cáo theo mẫu quy định\n5. Rà soát, phê duyệt\n6. Nộp báo cáo đúng hạn\n7. Lưu trữ bằng chứng nộp báo cáo",
     "• Regulatory Reporting AI: tự động tổng hợp dữ liệu và điền mẫu báo cáo\n• ETL + AI pipeline: trích xuất, chuyển đổi dữ liệu từ core banking\n• Obligation management: theo dõi hạn nộp báo cáo\n• RPA (Robotic Process Automation): tự động hóa nộp báo cáo lên cổng NHNN"),

    ("", "7.3",
     "Báo cáo nội bộ / Dashboard pháp chế",
     "Lập báo cáo định kỳ cho ban lãnh đạo về tình hình pháp lý: số vụ kiện, trạng thái tuân thủ, rủi ro pháp lý, tiến độ xử lý nợ...",
     "1. Xác định KPI/chỉ số cần báo cáo\n2. Thu thập dữ liệu từ hệ thống quản lý vụ việc\n3. Phân tích, tổng hợp\n4. Trực quan hóa dữ liệu\n5. Soạn thảo narrative báo cáo\n6. Trình ban lãnh đạo\n7. Lưu trữ",
     "• BI Dashboard: Power BI, Tableau tích hợp AI insights\n• LLM report narrative generator: tự động viết phần diễn giải báo cáo\n• AI anomaly detection: phát hiện bất thường trong dữ liệu pháp lý\n• Natural Language Generation (NLG): chuyển data thành văn bản báo cáo"),

    # NHOM 8
    ("NHÓM 8: KIỂM SOÁT RỦI RO PHÁP LÝ", "8.1",
     "Đánh giá rủi ro pháp lý sản phẩm/dịch vụ mới",
     "Đánh giá rủi ro pháp lý khi ngân hàng ra mắt sản phẩm, dịch vụ mới (fintech, open banking, cho vay số, bancassurance...) hoặc thay đổi quy trình nghiệp vụ.",
     "1. Nhận mô tả sản phẩm/dịch vụ mới\n2. Xác định các rủi ro pháp lý tiềm ẩn\n3. Tra cứu quy định áp dụng\n4. Đánh giá mức độ rủi ro (likelihood × impact)\n5. Đề xuất biện pháp kiểm soát/giảm thiểu\n6. Lập báo cáo đánh giá rủi ro pháp lý\n7. Theo dõi sau triển khai",
     "• AI Legal Risk Assessment: phân tích rủi ro pháp lý theo framework\n• LLM scenario analysis: mô phỏng kịch bản rủi ro pháp lý\n• RegTech AI: tự động mapping sản phẩm với quy định hiện hành\n• AI regulatory horizon scanning: phát hiện quy định sắp ban hành ảnh hưởng đến sản phẩm"),

    ("", "8.2",
     "Quản lý rủi ro tuân thủ (Compliance Risk Management)",
     "Xây dựng, vận hành ma trận rủi ro tuân thủ; theo dõi các rủi ro pháp lý trọng yếu; báo cáo cho Ủy ban Quản lý Rủi ro và HĐQT.",
     "1. Xây dựng compliance risk register\n2. Đánh giá định kỳ từng rủi ro\n3. Cập nhật mức độ rủi ro khi có thay đổi\n4. Thiết kế và kiểm tra hiệu quả kiểm soát\n5. Báo cáo lên Ủy ban Rủi ro\n6. Điều chỉnh chiến lược kiểm soát",
     "• GRC Platform AI: RSA Archer, ServiceNow IRM, 6clicks\n• AI risk scoring: chấm điểm rủi ro tự động\n• Predictive compliance analytics: dự báo rủi ro tuân thủ\n• Automated control testing: kiểm tra tự động hiệu quả kiểm soát"),
]

# ===== BUILD TABLE =====
# Columns: STT | Loại công việc | Mô tả chi tiết | Quy trình/Các bước | AI/Công nghệ
headers = ["STT", "Loại công việc", "Mô tả chi tiết công việc", "Quy trình / Các bước thực hiện", "Loại AI / Công nghệ hỗ trợ"]
col_widths = [Cm(1.0), Cm(3.5), Cm(5.5), Cm(7.5), Cm(8.0)]

GROUP_COLORS = {
    "NHÓM 1": "2E75B6",
    "NHÓM 2": "70AD47",
    "NHÓM 3": "ED7D31",
    "NHÓM 4": "C00000",
    "NHÓM 5": "7030A0",
    "NHÓM 6": "00B0F0",
    "NHÓM 7": "FF0066",
    "NHÓM 8": "833C00",
}

ROW_COLORS_LIGHT = {
    "NHÓM 1": "D6E4F0",
    "NHÓM 2": "E2EFDA",
    "NHÓM 3": "FCE4D6",
    "NHÓM 4": "FFDCDC",
    "NHÓM 5": "EAD1F5",
    "NHÓM 6": "DDEBF7",
    "NHÓM 7": "FFE0EE",
    "NHÓM 8": "F4CCCC",
}

# Count total rows needed
# For each group: 1 group header row + N data rows
total_rows = 1  # header
current_group = None
for row in data:
    if row[0] != "":
        total_rows += 1  # group header
    total_rows += 1  # data row

table = doc.add_table(rows=total_rows, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for row in table.rows:
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]

# Write header row
for j, h in enumerate(headers):
    add_header_cell(table.rows[0].cells[j], h, font_size=9, bg_color="1F4E79")

current_table_row = 1
current_group = None

for entry in data:
    nhom, stt, loai_cv, mo_ta, quy_trinh, ai_tech = entry

    # Determine group key
    if nhom != "":
        current_group = nhom[:6].strip()  # e.g. "NHÓM 1"
        group_color = GROUP_COLORS.get(current_group, "404040")
        # Add group header row (merged across all 5 cols)
        row = table.rows[current_table_row]
        cell = row.cells[0]
        # Merge all 5 cells
        merged = cell.merge(row.cells[4])
        set_cell_bg(merged, group_color)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = merged.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("  " + nhom)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Times New Roman'
        current_table_row += 1

    light_color = ROW_COLORS_LIGHT.get(current_group, "F2F2F2")
    row = table.rows[current_table_row]

    add_body_cell(row.cells[0], stt, bold=True, bg_color=light_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body_cell(row.cells[1], loai_cv, bold=True, bg_color=light_color)
    add_body_cell(row.cells[2], mo_ta, bg_color="FFFFFF")
    add_body_cell(row.cells[3], quy_trinh, bg_color="FFFFFF")
    add_body_cell(row.cells[4], ai_tech, bg_color="FFFFFF")

    current_table_row += 1

# ===== FOOTER NOTE =====
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_note = note.add_run(
    "Ghi chú: Bảng tổng hợp trên bao gồm 8 nhóm công việc chính với 24 loại công việc cụ thể của bộ phận Pháp chế ngân hàng. "
    "Các công nghệ AI được đề xuất bao gồm cả giải pháp thương mại quốc tế (Harvey AI, Luminance, Kira...) và hướng xây dựng "
    "giải pháp AI nội bộ (RAG chatbot, LLM tích hợp hệ thống ngân hàng). Mức độ ưu tiên triển khai nên dựa trên ROI, "
    "mức độ lặp lại của công việc và nguồn lực hiện có của Techcombank."
)
run_note.italic = True
run_note.font.size = Pt(8.5)
run_note.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
run_note.font.name = 'Times New Roman'

# ===== SAVE =====
output_path = "/workspace/UngDungAI_PhapChe_Techcombank.docx"
doc.save(output_path)
print(f"File saved: {output_path}")
