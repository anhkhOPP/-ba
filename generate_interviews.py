# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = "/workspace/phong_van_AI_phap_che"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold_prefix=None):
    """Add paragraph, optionally with a bold label prefix like 'PVV: '"""
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
    p.add_run(text)
    return p


def add_qa(doc, question, answer, interviewer="PVV", interviewee="NTL"):
    add_paragraph(doc, question, bold_prefix=f"{interviewer}: ")
    add_paragraph(doc, answer, bold_prefix=f"{interviewee}: ")
    doc.add_paragraph()


# =============================================================================
# PHỎNG VẤN 1: Chuyên gia pháp lý tại công ty Fintech
# =============================================================================
def create_interview_1():
    doc = Document()
    set_heading(doc, "BIÊN BẢN PHỎNG VẤN SÂU SỐ 01")
    set_heading(doc, "Chủ đề: Ứng dụng AI vào công tác pháp chế doanh nghiệp", level=2)

    doc.add_paragraph("Thông tin phỏng vấn:", style="Intense Quote")
    doc.add_paragraph("- Đối tượng phỏng vấn (NTL): Ông Nguyễn Minh Trí – Trưởng phòng Pháp chế, Công ty CP Tài chính Số Việt (FinTech)")
    doc.add_paragraph("- Người phỏng vấn (PVV): Nghiên cứu viên Lê Thị Hương")
    doc.add_paragraph("- Thời gian: 14h00 – 15h30, ngày 15 tháng 02 năm 2025")
    doc.add_paragraph("- Địa điểm: Phòng họp tầng 8, tòa nhà FPT Tower, Hà Nội")
    doc.add_paragraph("- Hình thức: Phỏng vấn trực tiếp, có ghi âm với sự đồng ý của đối tượng phỏng vấn")
    doc.add_paragraph()

    set_heading(doc, "NỘI DUNG PHỎNG VẤN", level=2)

    add_qa(doc,
        "Xin chào ông Trí, cảm ơn ông đã dành thời gian cho buổi phỏng vấn hôm nay. Trước tiên, ông có thể giới thiệu sơ lược về vai trò của mình tại công ty và bộ phận pháp chế không ạ?",
        "Xin chào chị Hương. Tôi làm Trưởng phòng Pháp chế tại Công ty CP Tài chính Số Việt được gần 5 năm nay. Công ty chúng tôi hoạt động trong lĩnh vực fintech – cho vay ngang hàng, ví điện tử và các dịch vụ tài chính số. Bộ phận pháp chế của tôi có 7 người, chúng tôi chịu trách nhiệm soạn thảo và rà soát hợp đồng, tuân thủ pháp luật, xử lý tranh chấp và tư vấn nội bộ. Đặc thù của fintech là môi trường pháp lý thay đổi rất nhanh, nên khối lượng công việc luôn rất lớn."
    )

    add_qa(doc,
        "Công ty ông đã bắt đầu ứng dụng AI vào công tác pháp chế từ khi nào và xuất phát từ nhu cầu gì?",
        "Chúng tôi bắt đầu thử nghiệm từ cuối năm 2022, nhưng triển khai chính thức vào đầu năm 2023. Nhu cầu xuất phát từ thực tế là bộ phận pháp chế của chúng tôi liên tục bị quá tải. Mỗi tháng chúng tôi phải rà soát hàng trăm hợp đồng đối tác, hàng chục điều khoản điều kiện sử dụng dịch vụ, chưa kể việc theo dõi văn bản pháp luật mới. Ban đầu chúng tôi dùng các công cụ AI đơn giản để tóm tắt văn bản. Sau đó chúng tôi đầu tư vào một nền tảng AI pháp lý chuyên biệt có khả năng phân tích hợp đồng, nhận diện rủi ro, và so sánh điều khoản với mẫu chuẩn của chúng tôi."
    )

    add_qa(doc,
        "Cụ thể AI đang hỗ trợ bộ phận của ông như thế nào trong công việc hàng ngày?",
        "Có ba ứng dụng chính. Thứ nhất là rà soát hợp đồng tự động – AI quét toàn bộ hợp đồng, đánh dấu các điều khoản bất lợi, thiếu sót hoặc mơ hồ, và so sánh với thư viện điều khoản chuẩn của chúng tôi. Thứ hai là tra cứu pháp luật – chúng tôi có một chatbot AI được huấn luyện trên cơ sở dữ liệu pháp luật Việt Nam, có thể trả lời câu hỏi pháp lý nội bộ trong vài giây. Thứ ba là theo dõi sự tuân thủ – AI tự động cập nhật khi có văn bản pháp luật mới và cảnh báo nếu có nội dung liên quan đến hoạt động kinh doanh của chúng tôi. Nhờ vậy, nhóm pháp chế của tôi tiết kiệm được khoảng 40% thời gian làm việc thủ công."
    )

    add_qa(doc,
        "Ông đánh giá thế nào về độ chính xác và độ tin cậy của các công cụ AI này?",
        "Tôi muốn thành thật rằng AI không hoàn hảo. Trong những tháng đầu triển khai, chúng tôi gặp khá nhiều trường hợp AI hiểu sai ngữ cảnh pháp lý, đặc biệt với các điều khoản đặc thù của luật Việt Nam mà không có nhiều dữ liệu huấn luyện. Tuy nhiên, sau khi chúng tôi bổ sung thêm dữ liệu đặc thù ngành và tinh chỉnh mô hình, độ chính xác đã cải thiện đáng kể – tôi ước tính khoảng 85-90% cho việc nhận diện rủi ro cơ bản. Quan trọng là chúng tôi luôn có luật sư con người xem xét lại kết quả của AI trước khi ra quyết định. AI là công cụ hỗ trợ, không phải người ra quyết định."
    )

    add_qa(doc,
        "Đâu là những thách thức lớn nhất mà ông gặp phải khi triển khai AI trong công tác pháp chế?",
        "Thách thức đầu tiên là về dữ liệu. Dữ liệu pháp lý bằng tiếng Việt còn rất hạn chế, và nhiều văn bản pháp luật Việt Nam chưa được số hóa đầy đủ. Thứ hai là vấn đề bảo mật – chúng tôi không thể đưa hợp đồng bảo mật lên các dịch vụ AI đám mây công cộng, nên phải đầu tư vào hạ tầng AI tại chỗ, chi phí khá cao. Thứ ba là yếu tố con người – ban đầu một số thành viên trong nhóm lo ngại AI sẽ thay thế việc làm của họ, tạo ra sự kháng cự. Chúng tôi phải mất nhiều thời gian để thay đổi tư duy, giúp họ hiểu rằng AI giải phóng họ khỏi công việc lặp lại để tập trung vào những vấn đề phức tạp hơn."
    )

    add_qa(doc,
        "Ông có ý kiến gì về khung pháp lý hiện tại liên quan đến việc sử dụng AI trong lĩnh vực pháp lý tại Việt Nam không?",
        "Đây là một khoảng trống đáng lo ngại. Hiện tại Việt Nam chưa có quy định cụ thể nào điều chỉnh việc sử dụng AI trong hành nghề pháp lý hay pháp chế doanh nghiệp. Chúng tôi hoạt động trong một vùng xám pháp lý. Về trách nhiệm pháp lý – nếu AI đưa ra kết quả sai và gây thiệt hại, trách nhiệm thuộc về ai? Về bảo mật dữ liệu – việc xử lý hợp đồng bằng AI phải tuân thủ Nghị định 13/2023 về bảo vệ dữ liệu cá nhân như thế nào? Tôi mong muốn các cơ quan quản lý sớm ban hành hướng dẫn cụ thể để doanh nghiệp có thể yên tâm đầu tư và triển khai."
    )

    add_qa(doc,
        "Nhìn về tương lai, ông thấy AI sẽ thay đổi công tác pháp chế doanh nghiệp như thế nào?",
        "Tôi tin rằng trong 5-10 năm tới, AI sẽ đảm nhiệm phần lớn công việc pháp chế mang tính lặp lại và tiêu chuẩn hóa – soạn thảo hợp đồng theo mẫu, rà soát tuân thủ, tra cứu pháp luật. Vai trò của luật sư và chuyên gia pháp chế sẽ dịch chuyển sang tư vấn chiến lược, đàm phán phức tạp, và những vấn đề đòi hỏi phán đoán đạo đức và kinh nghiệm thực tiễn. Tôi cũng kỳ vọng sẽ có các nền tảng AI pháp lý chuyên biệt cho thị trường Việt Nam, được xây dựng trên nền tảng pháp luật Việt Nam và được các cơ quan chức năng chứng nhận. Đó sẽ là bước tiến quan trọng để doanh nghiệp ứng dụng AI pháp chế một cách tự tin và đúng nghĩa."
    )

    add_qa(doc,
        "Cảm ơn ông rất nhiều về những chia sẻ vô cùng quý báu này. Ông có muốn bổ sung điều gì không?",
        "Tôi chỉ muốn nhắn nhủ rằng AI là công cụ rất mạnh mẽ nhưng không phải là đũa thần. Doanh nghiệp cần đầu tư nghiêm túc vào việc lựa chọn công cụ phù hợp, đào tạo nhân sự, và thiết lập quy trình kiểm soát chất lượng. Quan trọng hơn là không bao giờ để AI ra quyết định pháp lý mà không có sự giám sát của con người. Pháp lý là lĩnh vực có hậu quả nghiêm trọng nếu sai sót, nên ngưỡng kiểm soát phải rất cao. Cảm ơn chị đã có buổi phỏng vấn thú vị này."
    )

    doc.add_paragraph("--- HẾT PHỎNG VẤN ---")
    doc.save(os.path.join(OUTPUT_DIR, "PV01_Nguyen_Minh_Tri_Fintech.docx"))
    print("Đã tạo: PV01_Nguyen_Minh_Tri_Fintech.docx")


# =============================================================================
# PHỎNG VẤN 2: Chuyên gia pháp lý tại công ty sản xuất lớn
# =============================================================================
def create_interview_2():
    doc = Document()
    set_heading(doc, "BIÊN BẢN PHỎNG VẤN SÂU SỐ 02")
    set_heading(doc, "Chủ đề: Ứng dụng AI vào công tác pháp chế trong doanh nghiệp sản xuất", level=2)

    doc.add_paragraph("Thông tin phỏng vấn:", style="Intense Quote")
    doc.add_paragraph("- Đối tượng phỏng vấn (NTL): Bà Trần Thị Bích Ngọc – Giám đốc Pháp lý, Tập đoàn Sản xuất và Thương mại Á Đông")
    doc.add_paragraph("- Người phỏng vấn (PVV): Nghiên cứu viên Lê Thị Hương")
    doc.add_paragraph("- Thời gian: 09h00 – 10h45, ngày 22 tháng 02 năm 2025")
    doc.add_paragraph("- Địa điểm: Trụ sở Tập đoàn Á Đông, Khu công nghiệp Bình Dương")
    doc.add_paragraph("- Hình thức: Phỏng vấn trực tiếp, có ghi âm với sự đồng ý của đối tượng phỏng vấn")
    doc.add_paragraph()

    set_heading(doc, "NỘI DUNG PHỎNG VẤN", level=2)

    add_qa(doc,
        "Chào bà Ngọc, bà có thể chia sẻ về quy mô hoạt động của bộ phận pháp lý tại Á Đông và những áp lực công việc hiện tại không?",
        "Chào chị Hương. Tập đoàn Á Đông có hơn 15.000 nhân viên, 8 nhà máy sản xuất trải dài từ Bắc vào Nam và mạng lưới phân phối tại 12 quốc gia. Bộ phận pháp lý của tôi có 25 người, chia thành các nhóm: hợp đồng thương mại, lao động, sở hữu trí tuệ, tuân thủ quốc tế và tranh tụng. Áp lực lớn nhất là khối lượng hợp đồng thương mại – mỗi tháng chúng tôi xử lý hơn 500 hợp đồng ký kết với đối tác trong và ngoài nước, đặc biệt là hợp đồng song ngữ tiếng Việt – tiếng Anh. Đây là lý do chính khiến chúng tôi hướng đến AI như một giải pháp."
    )

    add_qa(doc,
        "Bà có thể mô tả cách tập đoàn đang triển khai AI trong công tác pháp chế không?",
        "Chúng tôi triển khai theo ba giai đoạn. Giai đoạn một từ 2022 đến 2023 là thử nghiệm với công cụ AI dịch thuật pháp lý – rất hữu ích cho hợp đồng quốc tế. Giai đoạn hai từ đầu 2024 là triển khai hệ thống quản lý hợp đồng tích hợp AI có thể tự động điền thông tin vào mẫu hợp đồng, phân loại hợp đồng theo mức độ rủi ro, và lên lịch nhắc nhở các điều khoản quan trọng như gia hạn, nghĩa vụ định kỳ. Giai đoạn ba đang trong quá trình triển khai là xây dựng hệ thống AI dự đoán tranh chấp – phân tích các hợp đồng hiện có để nhận diện sớm các điều khoản có nguy cơ dẫn đến xung đột."
    )

    add_qa(doc,
        "Những kết quả cụ thể nào bà có thể chia sẻ sau khi ứng dụng AI?",
        "Có một số con số tôi rất tự hào. Thời gian rà soát hợp đồng tiêu chuẩn giảm từ trung bình 4 giờ xuống còn 45 phút. Tỷ lệ phát hiện điều khoản bất lợi tăng từ 72% lên 91% so với khi làm thủ công – đây là kết quả kiểm tra nội bộ của chúng tôi. Chi phí thuê tư vấn pháp lý bên ngoài giảm 30% vì nhiều việc chúng tôi có thể tự xử lý nội bộ. Quan trọng hơn, nhóm pháp lý của tôi có thêm thời gian để tham gia vào các quyết định chiến lược của tập đoàn thay vì bị chôn vùi trong hồ sơ giấy tờ."
    )

    add_qa(doc,
        "Trong bối cảnh doanh nghiệp có hoạt động quốc tế, AI giúp ích gì cho việc tuân thủ đa thẩm quyền pháp lý?",
        "Đây là điểm mạnh nổi bật nhất. Chúng tôi phải tuân thủ luật pháp của 12 quốc gia đối tác, từ GDPR của châu Âu, đến luật thương mại Nhật Bản, Hàn Quốc, đến các quy định nhập khẩu của thị trường Mỹ. Hệ thống AI của chúng tôi được cập nhật liên tục với các thay đổi pháp luật tại các thị trường này và tự động cảnh báo khi một hợp đồng vi phạm hoặc có nguy cơ vi phạm quy định địa phương. Trước đây chúng tôi phải thuê công ty luật tại mỗi quốc gia để rà soát – rất tốn kém. Bây giờ AI làm công đoạn kiểm tra sơ bộ, chỉ những vấn đề phức tạp mới cần tư vấn nước ngoài."
    )

    add_qa(doc,
        "Bà nhận thấy những hạn chế nào của AI trong lĩnh vực pháp chế mà bà đang trực tiếp đối mặt?",
        "Hạn chế lớn nhất là AI không hiểu được ngữ cảnh kinh doanh. Đôi khi một điều khoản có vẻ bất lợi về mặt pháp lý thuần túy nhưng lại hợp lý trong bối cảnh đàm phán chiến lược – ví dụ chúng tôi chấp nhận điều khoản bồi thường cao hơn để đổi lấy giá tốt hơn hoặc ưu tiên giao hàng. AI không thể tự phán đoán điều này. Hạn chế thứ hai là AI gặp khó khăn với các tình huống pháp lý mới chưa có tiền lệ, đặc biệt là các vấn đề pháp lý xuyên biên giới mới phát sinh. Thứ ba là vấn đề cập nhật – pháp luật thay đổi liên tục và chúng tôi phải đảm bảo hệ thống AI luôn được cập nhật, đó là trách nhiệm của nhà cung cấp dịch vụ và cần hợp đồng bảo trì rõ ràng."
    )

    add_qa(doc,
        "Về vấn đề đào tạo nhân sự pháp chế để làm việc với AI, tập đoàn bà có những kinh nghiệm gì?",
        "Chúng tôi đã đầu tư rất nghiêm túc vào việc này. Trước hết là đào tạo kỹ năng sử dụng công cụ AI cơ bản – thực ra không khó, hầu hết nhân viên làm quen được trong 2-3 ngày. Phần quan trọng hơn là đào tạo tư duy làm việc với AI – cụ thể là kỹ năng kiểm tra và phê phán kết quả AI, nhận biết khi nào AI có thể sai, và kỹ năng prompt engineering để khai thác tốt nhất khả năng của công cụ. Chúng tôi cũng xây dựng một văn hóa 'AI là trợ lý, luật sư là người quyết định' – mọi quyết định pháp lý quan trọng đều phải có chữ ký của luật sư chịu trách nhiệm, AI chỉ cung cấp dữ liệu đầu vào."
    )

    add_qa(doc,
        "Bà có lời khuyên gì cho các doanh nghiệp đang cân nhắc đầu tư vào AI pháp chế?",
        "Lời khuyên đầu tiên là hãy bắt đầu từ vấn đề, không bắt đầu từ công nghệ. Xác định rõ điểm đau của bộ phận pháp chế của bạn là gì trước khi mua bất kỳ công cụ AI nào. Thứ hai, đừng kỳ vọng quá cao vào ngày đầu – cần ít nhất 6 tháng để hệ thống AI được tinh chỉnh phù hợp với đặc thù của doanh nghiệp. Thứ ba, đầu tư vào bảo mật dữ liệu từ đầu – đây không phải là chi phí tùy chọn. Và cuối cùng, hãy chọn nhà cung cấp có cam kết hỗ trợ lâu dài và cập nhật pháp luật thường xuyên – đây là yếu tố sống còn trong lĩnh vực pháp chế. Cảm ơn chị đã có buổi phỏng vấn rất thú vị."
    )

    doc.add_paragraph("--- HẾT PHỎNG VẤN ---")
    doc.save(os.path.join(OUTPUT_DIR, "PV02_Tran_Bich_Ngoc_SanXuat.docx"))
    print("Đã tạo: PV02_Tran_Bich_Ngoc_SanXuat.docx")


# =============================================================================
# PHỎNG VẤN 3: Học giả / Giảng viên luật
# =============================================================================
def create_interview_3():
    doc = Document()
    set_heading(doc, "BIÊN BẢN PHỎNG VẤN SÂU SỐ 03")
    set_heading(doc, "Chủ đề: Góc nhìn học thuật về AI trong pháp chế và hành nghề pháp lý", level=2)

    doc.add_paragraph("Thông tin phỏng vấn:", style="Intense Quote")
    doc.add_paragraph("- Đối tượng phỏng vấn (NTL): PGS.TS. Phạm Văn Hùng – Trưởng Bộ môn Luật Thương mại, Trường Đại học Luật Hà Nội")
    doc.add_paragraph("- Người phỏng vấn (PVV): Nghiên cứu viên Lê Thị Hương")
    doc.add_paragraph("- Thời gian: 10h00 – 11h30, ngày 05 tháng 03 năm 2025")
    doc.add_paragraph("- Địa điểm: Phòng làm việc của PGS.TS. Hùng, Trường Đại học Luật Hà Nội")
    doc.add_paragraph("- Hình thức: Phỏng vấn trực tiếp, có ghi âm với sự đồng ý của đối tượng phỏng vấn")
    doc.add_paragraph()

    set_heading(doc, "NỘI DUNG PHỎNG VẤN", level=2)

    add_qa(doc,
        "Kính chào PGS.TS. Hùng. Từ góc độ học thuật, ông nhìn nhận xu hướng ứng dụng AI vào lĩnh vực pháp lý như thế nào?",
        "Chào chị Hương. Đây là một xu hướng không thể đảo ngược và đang diễn ra với tốc độ rất nhanh trên toàn cầu. Từ góc độ lý thuyết, tôi phân loại ứng dụng AI trong pháp lý thành ba tầng: tầng thấp nhất là tự động hóa tác vụ đơn giản như soạn thảo theo mẫu, tìm kiếm văn bản; tầng giữa là phân tích và nhận diện mẫu như rà soát hợp đồng, phân tích án lệ; và tầng cao nhất – vẫn còn đang phát triển – là lý luận pháp lý, dự đoán kết quả tranh tụng. Ở Việt Nam hiện nay, chúng ta chủ yếu đang ở tầng thấp đến giữa. Tầng cao đòi hỏi sự hiểu biết sâu về hệ thống pháp luật, văn hóa pháp lý và thực tiễn xét xử của từng quốc gia."
    )

    add_qa(doc,
        "Ông đánh giá thế nào về tác động của AI đối với nghề luật và công tác pháp chế doanh nghiệp tại Việt Nam?",
        "Tác động là rất đa chiều. Về mặt tích cực, AI giúp phổ cập hóa dịch vụ pháp lý – trước đây chỉ doanh nghiệp lớn mới đủ khả năng duy trì đội ngũ pháp chế mạnh, nay với AI, doanh nghiệp vừa và nhỏ cũng có thể tiếp cận dịch vụ pháp lý chất lượng tốt hơn với chi phí thấp hơn. AI cũng giúp giảm thiểu sai sót do con người gây ra trong công việc lặp lại. Tuy nhiên, về mặt tiêu cực, có nguy cơ phụ thuộc thái quá vào AI dẫn đến giảm sút năng lực tư duy pháp lý độc lập của đội ngũ pháp chế. Cũng có những vấn đề về trách nhiệm pháp lý khi AI đưa ra tư vấn sai – hiện tại luật Việt Nam chưa có cơ chế quy trách nhiệm rõ ràng trong trường hợp này."
    )

    add_qa(doc,
        "Theo ông, những vấn đề pháp lý nào phát sinh từ việc sử dụng AI trong pháp chế cần được giải quyết khẩn cấp nhất?",
        "Tôi nhận diện ba vấn đề pháp lý cấp bách nhất. Thứ nhất là vấn đề tư cách pháp lý của quyết định AI – khi AI tham gia tư vấn pháp lý, liệu đó có cấu thành hành nghề luật không phép không? Điều này liên quan trực tiếp đến Luật Luật sư năm 2006 sửa đổi 2012 và quy định về hành nghề pháp lý. Thứ hai là vấn đề bảo mật thông tin khách hàng – khi dữ liệu hồ sơ pháp lý được đưa vào hệ thống AI, nghĩa vụ bảo mật thông tin theo quy tắc đạo đức nghề luật có bị vi phạm không? Thứ ba là vấn đề thiên vị thuật toán – nếu AI được huấn luyện trên dữ liệu án lệ có chứa thiên kiến lịch sử, nó có thể tái tạo và khuếch đại sự bất bình đẳng trong hệ thống pháp luật."
    )

    add_qa(doc,
        "Về mặt đạo đức pháp lý, sử dụng AI có gây ra những xung đột nào không?",
        "Câu hỏi này chạm đến trung tâm của cuộc tranh luận học thuật hiện nay. Nguyên tắc đầu tiên của đạo đức pháp lý là trung thành với thân chủ và phán đoán độc lập của luật sư. Khi luật sư phụ thuộc vào kết quả của AI mà không kiểm tra kỹ, nguyên tắc này bị xâm phạm. Có một câu hỏi triết học sâu hơn: pháp luật không chỉ là tập hợp quy tắc kỹ thuật mà còn là biểu hiện của giá trị xã hội và công lý – liệu AI có thể hiểu và áp dụng đúng chiều sâu đó không? Theo quan điểm của tôi, AI rất giỏi tìm kiếm và phân loại, nhưng lý luận pháp lý đòi hỏi sự thấu hiểu về con người và bối cảnh xã hội mà AI hiện tại chưa thể thực sự làm được."
    )

    add_qa(doc,
        "Trường Đại học Luật có những thay đổi gì trong chương trình giảng dạy để chuẩn bị cho sinh viên làm việc trong môi trường có AI không?",
        "Chúng tôi đang trong quá trình cải cách chương trình khá toàn diện. Từ năm học 2024-2025, chúng tôi đưa vào môn học 'Pháp lý và Công nghệ' bắt buộc cho sinh viên năm ba. Môn học này bao gồm các chủ đề về pháp luật điều chỉnh AI, sử dụng công cụ AI pháp lý trong thực hành, và đạo đức nghề nghiệp trong bối cảnh số hóa. Chúng tôi cũng tích hợp kỹ năng sử dụng cơ sở dữ liệu pháp lý và các công cụ tìm kiếm AI vào các môn học thực hành. Tuy nhiên, thách thức lớn nhất là đội ngũ giảng viên – chúng tôi cần đào tạo lại hoặc tuyển mới giảng viên có kiến thức liên ngành pháp lý và công nghệ, và đây là nguồn nhân lực rất khan hiếm tại Việt Nam."
    )

    add_qa(doc,
        "Ông có kiến nghị gì với các cơ quan nhà nước về việc xây dựng khung pháp lý cho AI trong lĩnh vực pháp lý không?",
        "Tôi có một số kiến nghị cụ thể. Về ngắn hạn, Bộ Tư pháp và Liên đoàn Luật sư Việt Nam cần sớm ban hành hướng dẫn về việc sử dụng AI trong hành nghề pháp lý, tương tự như hướng dẫn mà Law Society của Anh hay Bar Association của Mỹ đã ban hành. Về trung hạn, cần có quy định cụ thể về tiêu chuẩn kỹ thuật và kiểm định đối với phần mềm AI pháp lý – tương tự như chứng nhận phần mềm kế toán. Về dài hạn, chúng ta cần nghiên cứu để bổ sung vào Luật Luật sư và các văn bản pháp luật liên quan các quy định về quyền và trách nhiệm khi sử dụng AI trong hành nghề. Đặc biệt quan trọng là xây dựng cơ chế giải quyết tranh chấp và phân bổ trách nhiệm khi AI gây ra thiệt hại pháp lý."
    )

    add_qa(doc,
        "Cuối cùng, ông nhìn nhận Việt Nam đang ở đâu so với xu thế toàn cầu về AI pháp lý?",
        "Thẳng thắn mà nói, chúng ta đang đi sau các nước phát triển khoảng 3-5 năm về mặt ứng dụng thực tiễn, và khoảng 5-7 năm về mặt khung pháp lý điều chỉnh. Tuy nhiên, đây không hẳn là bất lợi hoàn toàn. Chúng ta có cơ hội học hỏi kinh nghiệm và tránh những sai lầm mà các nước đi trước đã mắc phải. Điểm mạnh của Việt Nam là chúng ta có hệ thống pháp luật thành văn, có khả năng tiếp thu và điều chỉnh nhanh. Điều quan trọng là cần có ý chí chính trị và sự đầu tư nghiêm túc từ nhà nước và doanh nghiệp để bắt kịp xu thế. Cảm ơn chị đã có cuộc phỏng vấn rất giá trị này."
    )

    doc.add_paragraph("--- HẾT PHỎNG VẤN ---")
    doc.save(os.path.join(OUTPUT_DIR, "PV03_Pham_Van_Hung_HocGia.docx"))
    print("Đã tạo: PV03_Pham_Van_Hung_HocGia.docx")


# =============================================================================
# PHỎNG VẤN 4: Cán bộ pháp chế tại cơ quan nhà nước
# =============================================================================
def create_interview_4():
    doc = Document()
    set_heading(doc, "BIÊN BẢN PHỎNG VẤN SÂU SỐ 04")
    set_heading(doc, "Chủ đề: Ứng dụng AI vào công tác pháp chế trong cơ quan nhà nước", level=2)

    doc.add_paragraph("Thông tin phỏng vấn:", style="Intense Quote")
    doc.add_paragraph("- Đối tượng phỏng vấn (NTL): Ông Lê Quang Dũng – Phó Vụ trưởng Vụ Pháp chế, Bộ Công Thương")
    doc.add_paragraph("- Người phỏng vấn (PVV): Nghiên cứu viên Lê Thị Hương")
    doc.add_paragraph("- Thời gian: 14h30 – 16h00, ngày 10 tháng 03 năm 2025")
    doc.add_paragraph("- Địa điểm: Trụ sở Bộ Công Thương, 54 Hai Bà Trưng, Hà Nội")
    doc.add_paragraph("- Hình thức: Phỏng vấn trực tiếp, có ghi âm với sự đồng ý của đối tượng phỏng vấn")
    doc.add_paragraph()

    set_heading(doc, "NỘI DUNG PHỎNG VẤN", level=2)

    add_qa(doc,
        "Thưa ông Dũng, Bộ Công Thương đang ứng dụng AI trong công tác pháp chế như thế nào? Ông có thể chia sẻ tổng quan không?",
        "Chào chị Hương. Bộ Công Thương là một trong những Bộ tiên phong trong việc ứng dụng công nghệ vào công tác pháp chế ở cơ quan nhà nước. Hiện tại chúng tôi đang triển khai hệ thống cơ sở dữ liệu pháp luật thương mại tích hợp AI có tên LEGAL-CT, được kết nối với Cơ sở dữ liệu quốc gia về pháp luật của Bộ Tư pháp. Hệ thống này cho phép tự động phân loại, lập chỉ mục và tìm kiếm thông minh trong số hơn 5.000 văn bản quy phạm pháp luật thuộc lĩnh vực thương mại. Chúng tôi cũng đang thử nghiệm AI trong việc soạn thảo văn bản pháp luật – AI gợi ý các điều khoản dựa trên quy định hiện hành và thông lệ quốc tế."
    )

    add_qa(doc,
        "Trong công tác xây dựng văn bản pháp luật, AI hỗ trợ như thế nào để đảm bảo tính nhất quán và hệ thống của pháp luật?",
        "Đây là ứng dụng tôi đánh giá cao nhất. Khi soạn thảo một nghị định hoặc thông tư mới, một trong những thách thức lớn nhất là đảm bảo không mâu thuẫn với các văn bản hiện hành và tương thích với điều ước quốc tế mà Việt Nam là thành viên. Trước đây, tổ soạn thảo phải đọc thủ công hàng trăm văn bản liên quan – rất tốn thời gian và dễ bỏ sót. Nay AI có thể quét toàn bộ hệ thống văn bản pháp luật hiện hành và chỉ ra ngay những điểm có thể mâu thuẫn hoặc chồng lấn. Nó cũng so sánh với các cam kết WTO, các FTA thế hệ mới như CPTPP, EVFTA để cảnh báo những điều khoản có thể vi phạm nghĩa vụ quốc tế. Điều này giúp nâng cao đáng kể chất lượng văn bản pháp luật."
    )

    add_qa(doc,
        "Về công tác thẩm định và rà soát văn bản pháp luật, AI đóng vai trò như thế nào?",
        "Chúng tôi đang xây dựng quy trình thẩm định văn bản tích hợp AI theo hai bước. Bước đầu tiên là rà soát kỹ thuật tự động – AI kiểm tra cấu trúc văn bản, định dạng, viện dẫn pháp lý và logic nội tại. Bước thứ hai là phân tích tác động pháp lý – AI dự đoán những lĩnh vực và văn bản nào có thể bị ảnh hưởng bởi quy định mới. Tuy nhiên, tôi muốn nhấn mạnh rằng trong hệ thống nhà nước, chúng tôi áp dụng nguyên tắc 'AI hỗ trợ, con người quyết định' rất nghiêm ngặt. Mọi văn bản pháp luật đều phải được cán bộ pháp chế có thẩm quyền thẩm định và chịu trách nhiệm – AI chỉ là công cụ hỗ trợ trong quy trình này."
    )

    add_qa(doc,
        "Có những vấn đề gì đặc thù khi ứng dụng AI trong cơ quan nhà nước so với doanh nghiệp tư nhân không?",
        "Có rất nhiều điểm đặc thù. Thứ nhất là về bảo mật thông tin nhà nước – nhiều hồ sơ, dự thảo văn bản pháp luật của chúng tôi thuộc danh mục bí mật nhà nước, tuyệt đối không thể sử dụng dịch vụ AI đám mây. Mọi hệ thống AI phải được triển khai trên hạ tầng tự quản lý của nhà nước và tuân thủ nghiêm ngặt quy định bảo mật. Thứ hai là tính trách nhiệm giải trình – trong khu vực nhà nước, mọi quyết định đều phải có người chịu trách nhiệm cụ thể, rõ ràng trước pháp luật. Điều này tạo ra rào cản tâm lý lớn đối với việc sử dụng AI. Thứ ba là tính minh bạch trong quy trình – quy trình xây dựng và thẩm định văn bản pháp luật phải công khai và có thể kiểm tra, trong khi thuật toán AI thường là 'hộp đen'. Đây là những thách thức chúng tôi đang nghiên cứu giải pháp."
    )

    add_qa(doc,
        "Bộ Công Thương có kế hoạch ứng dụng AI trong việc giải quyết khiếu nại, tranh chấp thương mại không?",
        "Chúng tôi đang nghiên cứu mô hình này. Hiện tại chúng tôi đã triển khai thử nghiệm một hệ thống AI hỗ trợ tra cứu thủ tục giải quyết tranh chấp thương mại, giúp doanh nghiệp xác định được cơ quan có thẩm quyền giải quyết và hồ sơ cần chuẩn bị. Về dài hạn, chúng tôi đang nghiên cứu ODR – Online Dispute Resolution tích hợp AI – cho các tranh chấp thương mại nhỏ và vừa. Kinh nghiệm từ các nước như Trung Quốc với hệ thống Tòa án Internet Alibaba cho thấy AI có thể hỗ trợ giải quyết nhanh chóng các tranh chấp thương mại điện tử đơn giản. Tuy nhiên, ở Việt Nam, chúng ta cần nghiên cứu kỹ về tính hợp lệ pháp lý của quyết định do AI hỗ trợ trước khi triển khai chính thức."
    )

    add_qa(doc,
        "Ông nhìn nhận vai trò của nhà nước trong việc định hướng và quản lý việc ứng dụng AI vào pháp chế như thế nào?",
        "Nhà nước có ba vai trò thiết yếu. Vai trò thứ nhất là kiến tạo hành lang pháp lý – ban hành quy định rõ ràng về tiêu chuẩn AI pháp lý, điều kiện sử dụng, trách nhiệm pháp lý. Đây là nền tảng để doanh nghiệp và các tổ chức pháp lý yên tâm đầu tư. Vai trò thứ hai là xây dựng hạ tầng số dùng chung – xây dựng cơ sở dữ liệu pháp luật quốc gia chuẩn hóa, mở cho phát triển AI, đặc biệt là chuẩn hóa ngôn ngữ pháp lý tiếng Việt để AI có thể hiểu và xử lý chính xác hơn. Vai trò thứ ba là dẫn dắt và làm gương – các cơ quan nhà nước cần đi đầu trong ứng dụng AI pháp lý có trách nhiệm, trở thành mô hình tham chiếu cho khu vực tư nhân. Cảm ơn chị đã có buổi phỏng vấn bổ ích."
    )

    doc.add_paragraph("--- HẾT PHỎNG VẤN ---")
    doc.save(os.path.join(OUTPUT_DIR, "PV04_Le_Quang_Dung_BoCongThuong.docx"))
    print("Đã tạo: PV04_Le_Quang_Dung_BoCongThuong.docx")


# =============================================================================
# PHỎNG VẤN 5: Luật sư tại công ty luật tư nhân
# =============================================================================
def create_interview_5():
    doc = Document()
    set_heading(doc, "BIÊN BẢN PHỎNG VẤN SÂU SỐ 05")
    set_heading(doc, "Chủ đề: AI trong thực hành pháp lý tại công ty luật tư nhân", level=2)

    doc.add_paragraph("Thông tin phỏng vấn:", style="Intense Quote")
    doc.add_paragraph("- Đối tượng phỏng vấn (NTL): Luật sư Nguyễn Thị Thanh Hà – Quản lý cấp cao, Công ty Luật TNHH Horizon Legal")
    doc.add_paragraph("- Người phỏng vấn (PVV): Nghiên cứu viên Lê Thị Hương")
    doc.add_paragraph("- Thời gian: 16h00 – 17h30, ngày 18 tháng 03 năm 2025")
    doc.add_paragraph("- Địa điểm: Văn phòng Horizon Legal, Tầng 15, Tòa nhà Lotte Center, Hà Nội")
    doc.add_paragraph("- Hình thức: Phỏng vấn trực tiếp, có ghi âm với sự đồng ý của đối tượng phỏng vấn")
    doc.add_paragraph()

    set_heading(doc, "NỘI DUNG PHỎNG VẤN", level=2)

    add_qa(doc,
        "Chào luật sư Hà. Bà có thể chia sẻ cách Horizon Legal đang ứng dụng AI trong hoạt động thực hành pháp lý hàng ngày không?",
        "Chào chị Hương. Horizon Legal là một công ty luật tầm trung với khoảng 30 luật sư và 20 nhân viên pháp lý. Chúng tôi bắt đầu ứng dụng AI từ năm 2023 và hiện có một hệ sinh thái công cụ AI khá đa dạng. Chúng tôi sử dụng AI trong nghiên cứu pháp lý – thay vì tìm kiếm thủ công trong cơ sở dữ liệu án lệ và văn bản pháp luật, AI tổng hợp và trình bày các quy định liên quan theo chủ đề câu hỏi trong vài phút. Chúng tôi dùng AI để rà soát due diligence – quét hàng trăm tài liệu trong các thương vụ M&A để nhận diện rủi ro. Và chúng tôi dùng AI để soạn thảo văn bản pháp lý theo mẫu và tóm tắt hồ sơ cho khách hàng."
    )

    add_qa(doc,
        "Ứng dụng nào bà thấy mang lại hiệu quả cao nhất và có thể đo lường được?",
        "Rõ ràng nhất là trong công việc due diligence pháp lý cho các thương vụ M&A. Trước đây, một thương vụ due diligence trung bình tốn 3-4 tuần với một nhóm 5-6 luật sư làm việc toàn thời gian. Hiện tại, chúng tôi rút ngắn xuống còn 7-10 ngày với cùng nhóm người, vì AI xử lý phần đọc và sơ lọc tài liệu – công việc chiếm 60-70% thời gian trước đây. Luật sư chỉ tập trung vào phân tích những vấn đề AI đã nhận diện. Điều này không chỉ tiết kiệm thời gian mà còn nâng cao chất lượng – luật sư tập trung hơn, ít mệt mỏi hơn và do đó ít bỏ sót hơn. Khách hàng rất hài lòng vì tiến độ nhanh hơn và phí dịch vụ cũng giảm."
    )

    add_qa(doc,
        "Horizon Legal có dùng AI để tư vấn trực tiếp cho khách hàng hoặc dự đoán kết quả vụ kiện không?",
        "Chúng tôi thử nghiệm tính năng dự đoán kết quả vụ kiện nhưng phải thừa nhận rằng hiệu quả còn khá hạn chế trong điều kiện Việt Nam. Lý do chính là hệ thống án lệ của Việt Nam chưa phong phú và hệ thống hóa như ở Mỹ hay Anh – là nơi AI dự đoán tranh tụng hoạt động hiệu quả nhất. Án lệ của Việt Nam mới được công bố chính thức từ năm 2016, số lượng còn hạn chế và chưa được số hóa đầy đủ. Về tư vấn trực tiếp cho khách hàng qua AI, chúng tôi có một chatbot pháp lý trên website trả lời các câu hỏi cơ bản về dịch vụ và thủ tục, nhưng tuyệt đối không thay thế tư vấn của luật sư. Chúng tôi có tuyên bố rõ ràng trong giao diện rằng đây là thông tin chung, không phải tư vấn pháp lý chuyên nghiệp."
    )

    add_qa(doc,
        "Bà quản lý mối quan hệ với khách hàng về việc sử dụng AI như thế nào? Khách hàng có lo ngại gì không?",
        "Đây là một chủ đề nhạy cảm mà chúng tôi phải xử lý rất cẩn thận. Phần lớn khách hàng doanh nghiệp lớn, đặc biệt là các tập đoàn nước ngoài, rất ủng hộ việc chúng tôi dùng AI vì họ đã quen với điều này ở các thị trường khác. Tuy nhiên, một số khách hàng trong nước, đặc biệt là các cơ quan nhà nước và tổ chức bảo thủ hơn, có lo ngại về bảo mật dữ liệu – liệu thông tin hồ sơ của họ có bị lộ lọt khi đưa vào hệ thống AI không? Để giải quyết điều này, chúng tôi đã xây dựng chính sách AI rõ ràng, ký phụ lục bảo mật dữ liệu với khách hàng, và chỉ sử dụng AI trên hạ tầng bảo mật riêng không kết nối bên ngoài đối với hồ sơ bảo mật. Tính minh bạch với khách hàng là nguyên tắc không thể thỏa hiệp."
    )

    add_qa(doc,
        "Về góc độ cạnh tranh, AI đang thay đổi cấu trúc thị trường dịch vụ pháp lý như thế nào?",
        "AI đang tạo ra sự tái cấu trúc khá sâu sắc. Trước đây, quy mô công ty luật là lợi thế cạnh tranh lớn – công ty lớn có nhiều luật sư hơn, xử lý được nhiều hơn. Nay AI thu hẹp khoảng cách này. Một công ty luật vừa như chúng tôi với AI có thể cạnh tranh về tốc độ và giá cả với các hãng luật lớn hơn trong những mảng công việc nhất định. Điều này tốt cho cạnh tranh thị trường và tốt cho khách hàng. Tuy nhiên, AI cũng tạo ra một phân khúc mới: công ty AI pháp lý – tức là các startup công nghệ cung cấp dịch vụ pháp lý hoàn toàn tự động không có luật sư con người. Đây là thách thức pháp lý và cạnh tranh mà ngành luật cần phải đối mặt. Liệu đây có phải là hành nghề pháp lý không phép không? Tôi cho rằng các cơ quan quản lý cần sớm có câu trả lời."
    )

    add_qa(doc,
        "Bà thấy kỹ năng nào quan trọng nhất mà một luật sư hay chuyên gia pháp chế cần có trong thời đại AI?",
        "Tôi suy nghĩ nhiều về câu hỏi này, kể cả trong quá trình tuyển dụng gần đây. Kỹ năng đầu tiên là tư duy phản biện và phán đoán pháp lý – khả năng nhìn ra điều AI bỏ sót, đặt câu hỏi về kết quả AI cung cấp. Đây là kỹ năng con người mà AI chưa thể thay thế. Thứ hai là kỹ năng giao tiếp và quan hệ khách hàng – AI không thể thay thế sự tin tưởng, cảm thông và hiểu biết sâu sắc về nhu cầu cá nhân của khách hàng mà một luật sư giỏi xây dựng được. Thứ ba – và đây là điều mới – là kỹ năng làm việc với AI: biết cách thiết kế câu hỏi, kiểm tra kết quả, và tích hợp AI vào quy trình làm việc một cách thông minh. Luật sư không biết dùng AI sẽ bị thay thế không phải bởi AI mà bởi luật sư khác biết dùng AI tốt hơn. Cảm ơn chị đã có buổi phỏng vấn rất sâu sắc."
    )

    doc.add_paragraph("--- HẾT PHỎNG VẤN ---")
    doc.save(os.path.join(OUTPUT_DIR, "PV05_Nguyen_Thanh_Ha_LuatSu.docx"))
    print("Đã tạo: PV05_Nguyen_Thanh_Ha_LuatSu.docx")


if __name__ == "__main__":
    create_interview_1()
    create_interview_2()
    create_interview_3()
    create_interview_4()
    create_interview_5()
    print("\nHoàn tất! 5 file phỏng vấn đã được tạo tại:", OUTPUT_DIR)
