# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/workspace/phong_van_AI_phap_che"


def set_cell_bg(cell, hex_color):
    """Đặt màu nền cho ô bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_title(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_h3(doc, text, color=(0x37, 0x86, 0x4E)):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*color)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_quote(doc, text, source=""):
    p = doc.add_paragraph(style='Intense Quote')
    run = p.add_run(f'"{text}"')
    run.font.italic = True
    if source:
        src_run = p.add_run(f'\n— {source}')
        src_run.font.bold = True
        src_run.font.italic = False
        src_run.font.size = Pt(9)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_node_table(doc, rows):
    """rows = list of (node_path, color_hex)"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ["Cấp", "Tên Node", "Mã hiệu"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], "2E74B5")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for (level, name, code, color) in rows:
        row = table.add_row().cells
        row[0].text = level
        row[1].text = name
        row[2].text = code
        if color:
            set_cell_bg(row[0], color)
            set_cell_bg(row[1], color)
            set_cell_bg(row[2], color)
    doc.add_paragraph()
    return table


def create_codebook():
    doc = Document()

    # =========================================================
    # TRANG BÌA
    # =========================================================
    doc.add_paragraph()
    add_title(doc, "CODEBOOK – SÁCH MÃ HÓA")
    p = doc.add_paragraph("Dự án nghiên cứu:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("ỨNG DỤNG TRÍ TUỆ NHÂN TẠO (AI) VÀO CÔNG TÁC PHÁP CHẾ")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    doc.add_paragraph()
    info = [
        ("Phiên bản", "1.0 – Mã hóa lần đầu (First-cycle Coding)"),
        ("Dữ liệu", "5 biên bản phỏng vấn sâu (PV01 – PV05)"),
        ("Phần mềm", "NVivo 14"),
        ("Người soạn", "Nghiên cứu viên Lê Thị Hương"),
        ("Ngày", "Tháng 4 năm 2025"),
    ]
    for lbl, val in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{lbl}: ")
        r.bold = True
        p.add_run(val)

    doc.add_page_break()

    # =========================================================
    # PHẦN 1: HƯỚNG DẪN SỬ DỤNG CODEBOOK
    # =========================================================
    add_h2(doc, "PHẦN 1: HƯỚNG DẪN SỬ DỤNG CODEBOOK")
    add_body(doc,
        "Codebook này được xây dựng theo phương pháp mã hóa nội dung định hướng (Directed Content Analysis), "
        "kết hợp giữa mã hóa suy diễn (deductive) từ khung lý thuyết và mã hóa quy nạp (inductive) từ dữ liệu thực tế."
    )

    add_h3(doc, "1.1. Đơn vị mã hóa (Coding Unit)")
    add_body(doc,
        "Đơn vị mã hóa là đoạn trả lời của người được phỏng vấn (NTL), có thể là một câu, "
        "nhiều câu hoặc cả một đoạn văn. Không mã hóa câu hỏi của người phỏng vấn (PVV), "
        "trừ trường hợp câu hỏi phản ánh giả định cần phân tích."
    )
    add_body(doc,
        "Nguyên tắc lựa chọn đoạn mã hóa (Code Selection):"
    )
    items = [
        "Chọn đoạn văn bản đủ ngữ nghĩa – không cắt quá ngắn làm mất ngữ cảnh, không chọn quá dài gây khó quản lý.",
        "Một đoạn văn có thể được mã hóa bởi NHIỀU node cùng lúc (multiple coding) – đây là thực hành chuẩn trong NVivo.",
        "Ưu tiên chọn đoạn trả lời trực tiếp của NTL, không bao gồm câu hỏi của PVV trong selection trừ khi cần thiết cho ngữ cảnh.",
        "Mỗi lần kéo chọn (code selection) trong NVivo phải tương ứng đúng với định nghĩa node trong bảng dưới đây.",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    add_h3(doc, "1.2. Ký hiệu mã hiệu (Node Code)")
    add_body(doc,
        "Mã hiệu node theo quy ước: [NHÓM_CHÍNH].[SỐ THỨ TỰ].[SỐ CON] "
        "Ví dụ: UNG_DUNG.01 = Node chính thứ 01 trong nhóm Ứng dụng; THACH_THUC.02.01 = Node con thứ 01 của node thứ 02 trong nhóm Thách thức."
    )

    add_h3(doc, "1.3. Quy tắc ưu tiên khi mã hóa mơ hồ")
    items2 = [
        "Nếu đoạn văn vừa nói về 'lợi ích' vừa nói về 'thách thức' → mã hóa cả hai node.",
        "Nếu không chắc chắn giữa hai node → chọn node cụ thể hơn (node lá thay vì node gốc).",
        "Nếu một đoạn hoàn toàn không khớp với bất kỳ node nào → ghi chú vào Memo 'Cần xem lại' và tạo Free Node tạm thời.",
        "Câu thông tin cá nhân giới thiệu (tên, chức vụ, tổ chức) → không mã hóa vào node chủ đề.",
    ]
    for item in items2:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================
    # PHẦN 2: CÂY NODE (NODE TREE)
    # =========================================================
    add_h2(doc, "PHẦN 2: CẤU TRÚC CÂY NODE TỔNG QUAN")
    add_body(doc,
        "Toàn bộ codebook được tổ chức thành 7 nhóm node chính (parent nodes), "
        "mỗi nhóm chia thành các node con (child nodes). Tổng cộng: 7 node chính, 28 node con."
    )

    tree_rows = [
        ("NHÓM 1", "NHÓM: BỐI CẢNH & ĐỘNG LỰC ỨNG DỤNG AI",    "",                 "D6E4F7"),
        ("  Node con", "Áp lực & quá tải công việc pháp chế",    "BCLC.01",          "EBF3FC"),
        ("  Node con", "Nhu cầu số hóa & hiện đại hóa",          "BCLC.02",          "EBF3FC"),
        ("  Node con", "Lộ trình & giai đoạn triển khai",        "BCLC.03",          "EBF3FC"),
        ("  Node con", "Đặc thù ngành / lĩnh vực",               "BCLC.04",          "EBF3FC"),

        ("NHÓM 2", "NHÓM: CÁC ỨNG DỤNG AI CỤ THỂ",              "",                 "D5E8D4"),
        ("  Node con", "Rà soát & phân tích hợp đồng",           "UNG_DUNG.01",      "EBF5EA"),
        ("  Node con", "Tra cứu pháp luật thông minh",           "UNG_DUNG.02",      "EBF5EA"),
        ("  Node con", "Soạn thảo văn bản pháp lý",              "UNG_DUNG.03",      "EBF5EA"),
        ("  Node con", "Theo dõi tuân thủ & cảnh báo",           "UNG_DUNG.04",      "EBF5EA"),
        ("  Node con", "Due diligence & M&A",                    "UNG_DUNG.05",      "EBF5EA"),
        ("  Node con", "Dịch thuật pháp lý",                     "UNG_DUNG.06",      "EBF5EA"),
        ("  Node con", "Xây dựng & thẩm định VBPL",              "UNG_DUNG.07",      "EBF5EA"),
        ("  Node con", "Giải quyết tranh chấp & ODR",            "UNG_DUNG.08",      "EBF5EA"),

        ("NHÓM 3", "NHÓM: LỢI ÍCH & HIỆU QUẢ",                 "",                 "FFF2CC"),
        ("  Node con", "Tiết kiệm thời gian",                    "LOI_ICH.01",       "FFFAEB"),
        ("  Node con", "Nâng cao chất lượng & độ chính xác",     "LOI_ICH.02",       "FFFAEB"),
        ("  Node con", "Giảm chi phí",                           "LOI_ICH.03",       "FFFAEB"),
        ("  Node con", "Phổ cập hóa dịch vụ pháp lý",           "LOI_ICH.04",       "FFFAEB"),

        ("NHÓM 4", "NHÓM: THÁCH THỨC & HẠN CHẾ",                "",                 "FCE4D6"),
        ("  Node con", "Hạn chế dữ liệu tiếng Việt",            "THACH_THUC.01",    "FDF2EE"),
        ("  Node con", "Bảo mật & an ninh thông tin",            "THACH_THUC.02",    "FDF2EE"),
        ("  Node con", "AI không hiểu ngữ cảnh",                "THACH_THUC.03",    "FDF2EE"),
        ("  Node con", "Kháng cự & tâm lý nhân sự",             "THACH_THUC.04",    "FDF2EE"),
        ("  Node con", "Chi phí & hạ tầng",                     "THACH_THUC.05",    "FDF2EE"),

        ("NHÓM 5", "NHÓM: KHUNG PHÁP LÝ & CHÍNH SÁCH",         "",                 "E8D5F5"),
        ("  Node con", "Khoảng trống pháp lý tại Việt Nam",     "PHAP_LY.01",       "F5EDFB"),
        ("  Node con", "Trách nhiệm pháp lý của AI",            "PHAP_LY.02",       "F5EDFB"),
        ("  Node con", "Kiến nghị chính sách",                  "PHAP_LY.03",       "F5EDFB"),
        ("  Node con", "Tuân thủ đa thẩm quyền quốc tế",       "PHAP_LY.04",       "F5EDFB"),

        ("NHÓM 6", "NHÓM: ĐẠO ĐỨC & TRÁCH NHIỆM",             "",                 "F4CCCC"),
        ("  Node con", "Nguyên tắc AI hỗ trợ – con người quyết định", "DAO_DUC.01", "FDEAEA"),
        ("  Node con", "Thiên vị thuật toán",                   "DAO_DUC.02",       "FDEAEA"),
        ("  Node con", "Bảo mật thông tin thân chủ",            "DAO_DUC.03",       "FDEAEA"),
        ("  Node con", "Minh bạch với khách hàng",              "DAO_DUC.04",       "FDEAEA"),

        ("NHÓM 7", "NHÓM: TƯƠNG LAI NGHỀ PHÁP CHẾ",            "",                 "D9D9D9"),
        ("  Node con", "Thay đổi vai trò luật sư / pháp chế",  "TUONG_LAI.01",     "F2F2F2"),
        ("  Node con", "Kỹ năng cần thiết trong thời đại AI",   "TUONG_LAI.02",     "F2F2F2"),
        ("  Node con", "Đào tạo & cải cách giáo dục luật",     "TUONG_LAI.03",     "F2F2F2"),
        ("  Node con", "Cạnh tranh thị trường pháp lý",        "TUONG_LAI.04",     "F2F2F2"),
        ("  Node con", "So sánh Việt Nam với quốc tế",         "TUONG_LAI.05",     "F2F2F2"),
    ]

    add_node_table(doc, tree_rows)
    doc.add_page_break()

    # =========================================================
    # PHẦN 3: ĐỊNH NGHĨA CHI TIẾT TỪNG NODE
    # =========================================================
    add_h2(doc, "PHẦN 3: ĐỊNH NGHĨA CHI TIẾT TỪNG NODE VÀ HƯỚNG DẪN CODE SELECTION")

    # ---- NHÓM 1 ----
    add_h3(doc, "NHÓM 1: BỐI CẢNH & ĐỘNG LỰC ỨNG DỤNG AI", color=(0x1F, 0x4E, 0x79))
    add_body(doc, "Nhóm này mã hóa các đoạn văn bản giải thích TẠI SAO và NHƯ THẾ NÀO tổ chức bắt đầu ứng dụng AI. Chú ý nhóm này tập trung vào bối cảnh xuất phát điểm, không phải hiệu quả đã đạt được.")
    doc.add_paragraph()

    nodes_g1 = [
        (
            "BCLC.01 – Áp lực & quá tải công việc pháp chế",
            "Các đoạn mô tả khối lượng công việc quá lớn, thiếu nhân lực, áp lực về thời gian và nguồn lực của bộ phận pháp chế trước khi có AI. Đây thường là lý do thúc đẩy việc tìm kiếm giải pháp AI.",
            "Bao gồm: đề cập đến số lượng hợp đồng, số lượng nhân viên không đủ, thời gian xử lý kéo dài, làm thêm giờ, quá tải.",
            "Không bao gồm: đánh giá lợi ích sau khi có AI (→ chuyển sang LOI_ICH).",
            "Mỗi tháng chúng tôi phải rà soát hàng trăm hợp đồng đối tác... bộ phận pháp chế của chúng tôi liên tục bị quá tải.",
            "NTL_01 (PV01)"
        ),
        (
            "BCLC.02 – Nhu cầu số hóa & hiện đại hóa",
            "Các đoạn thể hiện định hướng chiến lược muốn hiện đại hóa quy trình pháp chế, đón đầu xu hướng, hoặc áp lực từ ban lãnh đạo về chuyển đổi số.",
            "Bao gồm: nhắc đến chuyển đổi số, chiến lược công nghệ, tầm nhìn lãnh đạo, áp lực cạnh tranh.",
            "Không bao gồm: mô tả kỹ thuật về cách AI hoạt động (→ UNG_DUNG).",
            "AI đang tạo ra sự tái cấu trúc khá sâu sắc... Trước đây quy mô công ty luật là lợi thế cạnh tranh lớn.",
            "NTL_05 (PV05)"
        ),
        (
            "BCLC.03 – Lộ trình & giai đoạn triển khai",
            "Các đoạn mô tả quá trình từng bước triển khai AI: thử nghiệm, pilot, triển khai chính thức, mở rộng. Chú ý đến thời gian, các mốc quan trọng và thứ tự ưu tiên.",
            "Bao gồm: đề cập đến giai đoạn 1, 2, 3; năm bắt đầu; thử nghiệm pilot; triển khai toàn diện.",
            "Không bao gồm: kết quả sau triển khai (→ LOI_ICH) hay khó khăn gặp phải (→ THACH_THUC).",
            "Chúng tôi triển khai theo ba giai đoạn. Giai đoạn một từ 2022 đến 2023 là thử nghiệm với công cụ AI dịch thuật pháp lý.",
            "NTL_02 (PV02)"
        ),
        (
            "BCLC.04 – Đặc thù ngành / lĩnh vực",
            "Các đoạn giải thích những đặc điểm riêng của ngành/lĩnh vực khiến nhu cầu AI trở nên đặc biệt quan trọng hoặc đặc biệt khó.",
            "Bao gồm: fintech thay đổi nhanh, quy định quốc tế phức tạp, bí mật nhà nước, đặc thù sản xuất.",
            "Không bao gồm: đặc thù tạo ra thách thức kỹ thuật (→ THACH_THUC).",
            "Đặc thù của fintech là môi trường pháp lý thay đổi rất nhanh, nên khối lượng công việc luôn rất lớn.",
            "NTL_01 (PV01)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g1:
        add_h3(doc, tieu_de, color=(0x2E, 0x74, 0xB5))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    # ---- NHÓM 2 ----
    add_h3(doc, "NHÓM 2: CÁC ỨNG DỤNG AI CỤ THỂ", color=(0x37, 0x86, 0x4E))
    add_body(doc, "Nhóm này mã hóa các đoạn mô tả CỤ THỂ AI đang làm GÌ trong công tác pháp chế. Đây là nhóm node quan trọng nhất và có nhiều node con nhất.")
    doc.add_paragraph()

    nodes_g2 = [
        (
            "UNG_DUNG.01 – Rà soát & phân tích hợp đồng",
            "Các đoạn mô tả việc AI tự động đọc, phân tích, đánh dấu rủi ro hoặc so sánh hợp đồng. Bao gồm cả việc AI nhận diện điều khoản bất lợi, thiếu sót hay mâu thuẫn.",
            "Bao gồm: contract review, nhận diện rủi ro hợp đồng, so sánh điều khoản, đánh dấu bất thường.",
            "Không bao gồm: soạn thảo hợp đồng mới (→ UNG_DUNG.03).",
            "AI quét toàn bộ hợp đồng, đánh dấu các điều khoản bất lợi, thiếu sót hoặc mơ hồ, và so sánh với thư viện điều khoản chuẩn.",
            "NTL_01 (PV01)"
        ),
        (
            "UNG_DUNG.02 – Tra cứu pháp luật thông minh",
            "Các đoạn mô tả AI hỗ trợ tìm kiếm, tóm tắt, giải đáp câu hỏi pháp lý từ cơ sở dữ liệu văn bản quy phạm pháp luật.",
            "Bao gồm: chatbot pháp lý, legal research AI, tìm kiếm văn bản quy phạm, trả lời câu hỏi nội bộ.",
            "Không bao gồm: cảnh báo văn bản mới (→ UNG_DUNG.04).",
            "Chúng tôi có một chatbot AI được huấn luyện trên cơ sở dữ liệu pháp luật Việt Nam, có thể trả lời câu hỏi pháp lý nội bộ trong vài giây.",
            "NTL_01 (PV01)"
        ),
        (
            "UNG_DUNG.03 – Soạn thảo văn bản pháp lý",
            "Các đoạn mô tả AI hỗ trợ tạo ra bản thảo hợp đồng, văn bản pháp lý theo mẫu, điền thông tin tự động.",
            "Bao gồm: auto-drafting, điền mẫu, gợi ý điều khoản, tóm tắt hồ sơ.",
            "Không bao gồm: rà soát hợp đồng đã có (→ UNG_DUNG.01), soạn thảo VBPL nhà nước (→ UNG_DUNG.07).",
            "Chúng tôi dùng AI để soạn thảo văn bản pháp lý theo mẫu và tóm tắt hồ sơ cho khách hàng.",
            "NTL_05 (PV05)"
        ),
        (
            "UNG_DUNG.04 – Theo dõi tuân thủ & cảnh báo",
            "Các đoạn mô tả AI theo dõi liên tục sự thay đổi pháp luật và cảnh báo khi có văn bản mới hoặc nguy cơ không tuân thủ.",
            "Bao gồm: compliance monitoring, cảnh báo văn bản mới, theo dõi nghĩa vụ định kỳ.",
            "Không bao gồm: tra cứu chủ động (→ UNG_DUNG.02).",
            "AI tự động cập nhật khi có văn bản pháp luật mới và cảnh báo nếu có nội dung liên quan đến hoạt động kinh doanh của chúng tôi.",
            "NTL_01 (PV01)"
        ),
        (
            "UNG_DUNG.05 – Due diligence & M&A",
            "Các đoạn mô tả AI xử lý khối lượng lớn tài liệu trong các thương vụ mua bán, sáp nhập hay đầu tư.",
            "Bao gồm: AI review tài liệu M&A, quét data room, sơ lọc hồ sơ.",
            "Không bao gồm: rà soát hợp đồng thông thường (→ UNG_DUNG.01).",
            "AI xử lý phần đọc và sơ lọc tài liệu – công việc chiếm 60-70% thời gian trước đây. Luật sư chỉ tập trung vào phân tích những vấn đề AI đã nhận diện.",
            "NTL_05 (PV05)"
        ),
        (
            "UNG_DUNG.06 – Dịch thuật pháp lý",
            "Các đoạn mô tả AI dịch hợp đồng, văn bản pháp lý giữa các ngôn ngữ, đặc biệt trong bối cảnh hợp đồng quốc tế.",
            "Bao gồm: AI dịch hợp đồng song ngữ, giải thích thuật ngữ pháp lý nước ngoài.",
            "Không bao gồm: tuân thủ pháp luật nước ngoài nói chung (→ PHAP_LY.04).",
            "Giai đoạn một từ 2022 đến 2023 là thử nghiệm với công cụ AI dịch thuật pháp lý – rất hữu ích cho hợp đồng quốc tế.",
            "NTL_02 (PV02)"
        ),
        (
            "UNG_DUNG.07 – Xây dựng & thẩm định văn bản pháp luật (VBPL)",
            "Các đoạn mô tả AI hỗ trợ soạn thảo, kiểm tra tính nhất quán, phát hiện mâu thuẫn trong quá trình làm luật tại cơ quan nhà nước.",
            "Bao gồm: AI rà soát VBPL, so sánh nghị định/thông tư, kiểm tra cam kết quốc tế.",
            "Không bao gồm: tra cứu VBPL đơn giản (→ UNG_DUNG.02).",
            "AI có thể quét toàn bộ hệ thống văn bản pháp luật hiện hành và chỉ ra ngay những điểm có thể mâu thuẫn hoặc chồng lấn.",
            "NTL_04 (PV04)"
        ),
        (
            "UNG_DUNG.08 – Giải quyết tranh chấp & ODR",
            "Các đoạn mô tả AI hỗ trợ giải quyết tranh chấp, dự đoán kết quả vụ kiện, hoặc nền tảng ODR (Online Dispute Resolution).",
            "Bao gồm: dự đoán kết quả tranh tụng, AI trong tòa án trực tuyến, hỗ trợ thủ tục khiếu nại.",
            "Không bao gồm: phân tích hợp đồng để ngăn ngừa tranh chấp (→ UNG_DUNG.01).",
            "Chúng tôi đang nghiên cứu ODR – Online Dispute Resolution tích hợp AI – cho các tranh chấp thương mại nhỏ và vừa.",
            "NTL_04 (PV04)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g2:
        add_h3(doc, tieu_de, color=(0x37, 0x86, 0x4E))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- NHÓM 3 ----
    add_h3(doc, "NHÓM 3: LỢI ÍCH & HIỆU QUẢ", color=(0xBF, 0x85, 0x00))
    add_body(doc, "Nhóm này mã hóa các đoạn nói về KẾT QUẢ ĐẠT ĐƯỢC sau khi ứng dụng AI – có thể là số liệu cụ thể hoặc nhận định định tính.")
    doc.add_paragraph()

    nodes_g3 = [
        (
            "LOI_ICH.01 – Tiết kiệm thời gian",
            "Các đoạn so sánh thời gian trước và sau khi có AI, hoặc nhận định AI giúp tăng tốc độ xử lý công việc.",
            "Bao gồm: số liệu thời gian cụ thể, so sánh trước/sau, giảm giờ làm thêm.",
            "Không bao gồm: tiết kiệm chi phí bằng tiền (→ LOI_ICH.03).",
            "Thời gian rà soát hợp đồng tiêu chuẩn giảm từ trung bình 4 giờ xuống còn 45 phút.",
            "NTL_02 (PV02)"
        ),
        (
            "LOI_ICH.02 – Nâng cao chất lượng & độ chính xác",
            "Các đoạn mô tả AI giúp phát hiện nhiều lỗi hơn, ít bỏ sót hơn, kết quả tốt hơn so với làm thủ công.",
            "Bao gồm: tỷ lệ phát hiện lỗi, độ chính xác, giảm sai sót.",
            "Không bao gồm: đánh giá hạn chế về độ chính xác (→ THACH_THUC.03).",
            "Tỷ lệ phát hiện điều khoản bất lợi tăng từ 72% lên 91% so với khi làm thủ công.",
            "NTL_02 (PV02)"
        ),
        (
            "LOI_ICH.03 – Giảm chi phí",
            "Các đoạn đề cập đến tiết kiệm chi phí tài chính: giảm phí tư vấn bên ngoài, giảm nhân lực cần thiết, ROI.",
            "Bao gồm: giảm phí thuê luật sư ngoài, giảm nhân sự, so sánh chi phí.",
            "Không bao gồm: tiết kiệm thời gian không gắn với tiền (→ LOI_ICH.01).",
            "Chi phí thuê tư vấn pháp lý bên ngoài giảm 30% vì nhiều việc chúng tôi có thể tự xử lý nội bộ.",
            "NTL_02 (PV02)"
        ),
        (
            "LOI_ICH.04 – Phổ cập hóa dịch vụ pháp lý",
            "Các đoạn nói về việc AI giúp SME, cơ quan nhỏ, người dân tiếp cận dịch vụ pháp lý tốt hơn với chi phí thấp hơn.",
            "Bao gồm: dân chủ hóa pháp lý, SME, bình đẳng tiếp cận.",
            "Không bao gồm: giảm chi phí cho doanh nghiệp lớn (→ LOI_ICH.03).",
            "AI giúp phổ cập hóa dịch vụ pháp lý – trước đây chỉ doanh nghiệp lớn mới đủ khả năng duy trì đội ngũ pháp chế mạnh, nay với AI, doanh nghiệp vừa và nhỏ cũng có thể tiếp cận dịch vụ pháp lý chất lượng tốt hơn.",
            "NTL_03 (PV03)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g3:
        add_h3(doc, tieu_de, color=(0xBF, 0x85, 0x00))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- NHÓM 4 ----
    add_h3(doc, "NHÓM 4: THÁCH THỨC & HẠN CHẾ", color=(0xC0, 0x50, 0x20))
    add_body(doc, "Nhóm này mã hóa các đoạn nói về KHÓ KHĂN, TRĂN TRỞ, VẤN ĐỀ chưa giải quyết được. Đây là nhóm node quan trọng để phân tích gap giữa kỳ vọng và thực tế.")
    doc.add_paragraph()

    nodes_g4 = [
        (
            "THACH_THUC.01 – Hạn chế dữ liệu tiếng Việt",
            "Các đoạn mô tả việc AI gặp khó khăn vì thiếu dữ liệu pháp lý chất lượng bằng tiếng Việt, pháp luật Việt Nam chưa số hóa đầy đủ.",
            "Bao gồm: thiếu corpus tiếng Việt, văn bản chưa số hóa, AI chưa hiểu luật Việt Nam.",
            "Không bao gồm: vấn đề bảo mật dữ liệu (→ THACH_THUC.02).",
            "Dữ liệu pháp lý bằng tiếng Việt còn rất hạn chế, và nhiều văn bản pháp luật Việt Nam chưa được số hóa đầy đủ.",
            "NTL_01 (PV01)"
        ),
        (
            "THACH_THUC.02 – Bảo mật & an ninh thông tin",
            "Các đoạn nói về lo ngại khi đưa tài liệu mật lên AI cloud, giải pháp on-premise, rủi ro lộ lọt thông tin.",
            "Bao gồm: AI cloud vs on-premise, dữ liệu bí mật, chính sách bảo mật, bí mật nhà nước.",
            "Không bao gồm: bảo mật thông tin thân chủ theo đạo đức nghề nghiệp (→ DAO_DUC.03).",
            "Chúng tôi không thể đưa hợp đồng bảo mật lên các dịch vụ AI đám mây công cộng, nên phải đầu tư vào hạ tầng AI tại chỗ, chi phí khá cao.",
            "NTL_01 (PV01)"
        ),
        (
            "THACH_THUC.03 – AI không hiểu ngữ cảnh",
            "Các đoạn mô tả AI đưa ra kết quả kỹ thuật đúng nhưng thiếu hiểu biết về bối cảnh kinh doanh, chiến lược, hoặc văn hóa pháp lý.",
            "Bao gồm: AI không hiểu ý đồ đàm phán, thiếu hiểu biết thực tiễn, sai sót do thiếu ngữ cảnh.",
            "Không bao gồm: thiếu dữ liệu huấn luyện (→ THACH_THUC.01).",
            "Hạn chế lớn nhất là AI không hiểu được ngữ cảnh kinh doanh. Đôi khi một điều khoản có vẻ bất lợi về mặt pháp lý thuần túy nhưng lại hợp lý trong bối cảnh đàm phán chiến lược.",
            "NTL_02 (PV02)"
        ),
        (
            "THACH_THUC.04 – Kháng cự & tâm lý nhân sự",
            "Các đoạn nói về lo ngại AI thay thế việc làm, sự kháng cự của nhân viên, thách thức thay đổi văn hóa tổ chức.",
            "Bao gồm: lo sợ mất việc, kháng cự thay đổi, cần thay đổi tư duy.",
            "Không bao gồm: đào tạo kỹ năng (→ TUONG_LAI.02, TUONG_LAI.03).",
            "Ban đầu một số thành viên trong nhóm lo ngại AI sẽ thay thế việc làm của họ, tạo ra sự kháng cự.",
            "NTL_01 (PV01)"
        ),
        (
            "THACH_THUC.05 – Chi phí & hạ tầng",
            "Các đoạn đề cập đến chi phí đầu tư, hạ tầng kỹ thuật cần thiết, rào cản tài chính khi triển khai AI.",
            "Bao gồm: chi phí mua/thuê phần mềm, máy chủ on-premise, chi phí tích hợp hệ thống.",
            "Không bao gồm: chi phí đào tạo nhân sự (→ TUONG_LAI.03).",
            "Phải đầu tư vào hạ tầng AI tại chỗ, chi phí khá cao.",
            "NTL_01 (PV01)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g4:
        add_h3(doc, tieu_de, color=(0xC0, 0x50, 0x20))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- NHÓM 5 ----
    add_h3(doc, "NHÓM 5: KHUNG PHÁP LÝ & CHÍNH SÁCH", color=(0x5B, 0x2C, 0x8D))
    add_body(doc, "Nhóm này mã hóa các đoạn bàn về luật pháp, quy định, chính sách NHÀ NƯỚC liên quan đến AI trong pháp lý – không phải AI hỗ trợ làm việc, mà là luật điều chỉnh AI.")
    doc.add_paragraph()

    nodes_g5 = [
        (
            "PHAP_LY.01 – Khoảng trống pháp lý tại Việt Nam",
            "Các đoạn nhận xét rằng pháp luật Việt Nam chưa có quy định đầy đủ hoặc rõ ràng về AI trong lĩnh vực pháp lý.",
            "Bao gồm: vùng xám pháp lý, thiếu quy định, Việt Nam chưa có...",
            "Không bao gồm: so sánh với nước ngoài (→ TUONG_LAI.05).",
            "Hiện tại Việt Nam chưa có quy định cụ thể nào điều chỉnh việc sử dụng AI trong hành nghề pháp lý hay pháp chế doanh nghiệp. Chúng tôi hoạt động trong một vùng xám pháp lý.",
            "NTL_01 (PV01)"
        ),
        (
            "PHAP_LY.02 – Trách nhiệm pháp lý của AI",
            "Các đoạn đặt câu hỏi hoặc thảo luận về ai chịu trách nhiệm khi AI gây ra sai sót hay thiệt hại pháp lý.",
            "Bao gồm: quy trách nhiệm, liability, lỗi của AI do ai chịu.",
            "Không bao gồm: trách nhiệm đạo đức (→ DAO_DUC.01).",
            "Về trách nhiệm pháp lý – nếu AI đưa ra kết quả sai và gây thiệt hại, trách nhiệm thuộc về ai?",
            "NTL_01 (PV01)"
        ),
        (
            "PHAP_LY.03 – Kiến nghị chính sách",
            "Các đoạn đề xuất cụ thể với cơ quan nhà nước về việc xây dựng quy định, tiêu chuẩn, hướng dẫn về AI pháp lý.",
            "Bao gồm: đề xuất Bộ, Quốc hội; kiến nghị ban hành; xây dựng tiêu chuẩn.",
            "Không bao gồm: nhận xét chung về khoảng trống pháp lý không kèm đề xuất (→ PHAP_LY.01).",
            "Bộ Tư pháp và Liên đoàn Luật sư Việt Nam cần sớm ban hành hướng dẫn về việc sử dụng AI trong hành nghề pháp lý.",
            "NTL_03 (PV03)"
        ),
        (
            "PHAP_LY.04 – Tuân thủ đa thẩm quyền quốc tế",
            "Các đoạn mô tả AI hỗ trợ tuân thủ đồng thời pháp luật của nhiều quốc gia, hoặc thách thức của việc tuân thủ nhiều hệ thống pháp luật.",
            "Bao gồm: GDPR, WTO, CPTPP, EVFTA, luật nước ngoài, multi-jurisdiction.",
            "Không bao gồm: dịch thuật văn bản (→ UNG_DUNG.06).",
            "Chúng tôi phải tuân thủ luật pháp của 12 quốc gia đối tác, từ GDPR của châu Âu, đến luật thương mại Nhật Bản, Hàn Quốc, đến các quy định nhập khẩu của thị trường Mỹ.",
            "NTL_02 (PV02)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g5:
        add_h3(doc, tieu_de, color=(0x5B, 0x2C, 0x8D))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    # ---- NHÓM 6 ----
    add_h3(doc, "NHÓM 6: ĐẠO ĐỨC & TRÁCH NHIỆM", color=(0x96, 0x23, 0x1C))
    add_body(doc, "Nhóm này mã hóa các đoạn liên quan đến các nguyên tắc đạo đức nghề nghiệp, giá trị, và trách nhiệm con người trong việc dùng AI.")
    doc.add_paragraph()

    nodes_g6 = [
        (
            "DAO_DUC.01 – Nguyên tắc AI hỗ trợ – con người quyết định",
            "Các đoạn khẳng định rõ rằng AI chỉ là công cụ, luật sư/pháp chế viên phải là người đưa ra quyết định cuối cùng.",
            "Bao gồm: AI là trợ lý, human in the loop, con người chịu trách nhiệm.",
            "Không bao gồm: trách nhiệm pháp lý khi AI sai (→ PHAP_LY.02).",
            "AI là công cụ hỗ trợ, không phải người ra quyết định. Mọi quyết định pháp lý quan trọng đều phải có chữ ký của luật sư chịu trách nhiệm.",
            "NTL_01 (PV01)"
        ),
        (
            "DAO_DUC.02 – Thiên vị thuật toán",
            "Các đoạn đề cập đến nguy cơ AI tái tạo hoặc khuếch đại sự bất công, thiên kiến có trong dữ liệu lịch sử.",
            "Bao gồm: bias, thiên kiến dữ liệu, bất bình đẳng trong hệ thống pháp luật.",
            "Không bao gồm: sai sót kỹ thuật thông thường của AI (→ THACH_THUC.03).",
            "Nếu AI được huấn luyện trên dữ liệu án lệ có chứa thiên kiến lịch sử, nó có thể tái tạo và khuếch đại sự bất bình đẳng trong hệ thống pháp luật.",
            "NTL_03 (PV03)"
        ),
        (
            "DAO_DUC.03 – Bảo mật thông tin thân chủ",
            "Các đoạn liên quan đến nghĩa vụ đạo đức nghề nghiệp của luật sư về bảo mật thông tin khi dùng AI.",
            "Bao gồm: attorney-client privilege, quy tắc đạo đức nghề luật, bảo mật hồ sơ thân chủ.",
            "Không bao gồm: bảo mật dữ liệu kỹ thuật (→ THACH_THUC.02).",
            "Khi dữ liệu hồ sơ pháp lý được đưa vào hệ thống AI, nghĩa vụ bảo mật thông tin theo quy tắc đạo đức nghề luật có bị vi phạm không?",
            "NTL_03 (PV03)"
        ),
        (
            "DAO_DUC.04 – Minh bạch với khách hàng",
            "Các đoạn mô tả việc thông báo cho khách hàng về việc sử dụng AI, chính sách AI, thỏa thuận minh bạch.",
            "Bao gồm: thông báo sử dụng AI, chính sách AI, phụ lục bảo mật dữ liệu với khách hàng.",
            "Không bao gồm: minh bạch về thuật toán với xã hội nói chung (→ PHAP_LY.01).",
            "Tính minh bạch với khách hàng là nguyên tắc không thể thỏa hiệp. Chúng tôi đã xây dựng chính sách AI rõ ràng, ký phụ lục bảo mật dữ liệu với khách hàng.",
            "NTL_05 (PV05)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g6:
        add_h3(doc, tieu_de, color=(0x96, 0x23, 0x1C))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    doc.add_page_break()

    # ---- NHÓM 7 ----
    add_h3(doc, "NHÓM 7: TƯƠNG LAI NGHỀ PHÁP CHẾ", color=(0x40, 0x40, 0x40))
    add_body(doc, "Nhóm này mã hóa các đoạn bàn về xu hướng TƯƠNG LAI: nghề pháp chế sẽ thay đổi như thế nào, kỹ năng nào cần có, đào tạo như thế nào.")
    doc.add_paragraph()

    nodes_g7 = [
        (
            "TUONG_LAI.01 – Thay đổi vai trò luật sư / pháp chế",
            "Các đoạn dự đoán hoặc mô tả sự dịch chuyển trong vai trò và chức năng của luật sư, pháp chế viên khi AI đảm nhiệm công việc lặp lại.",
            "Bao gồm: dịch chuyển sang tư vấn chiến lược, AI thay thế công việc routine.",
            "Không bao gồm: kỹ năng cụ thể cần có (→ TUONG_LAI.02).",
            "Vai trò của luật sư và chuyên gia pháp chế sẽ dịch chuyển sang tư vấn chiến lược, đàm phán phức tạp, và những vấn đề đòi hỏi phán đoán đạo đức và kinh nghiệm thực tiễn.",
            "NTL_01 (PV01)"
        ),
        (
            "TUONG_LAI.02 – Kỹ năng cần thiết trong thời đại AI",
            "Các đoạn liệt kê hoặc mô tả kỹ năng mà luật sư/pháp chế viên cần phát triển để làm việc hiệu quả với AI.",
            "Bao gồm: tư duy phản biện, kỹ năng prompt, kiểm tra kết quả AI, kỹ năng mềm.",
            "Không bao gồm: chương trình đào tạo cụ thể (→ TUONG_LAI.03).",
            "Luật sư không biết dùng AI sẽ bị thay thế không phải bởi AI mà bởi luật sư khác biết dùng AI tốt hơn.",
            "NTL_05 (PV05)"
        ),
        (
            "TUONG_LAI.03 – Đào tạo & cải cách giáo dục luật",
            "Các đoạn mô tả sự thay đổi trong chương trình đào tạo luật, khóa học mới, cải cách giáo dục để đáp ứng thời đại AI.",
            "Bao gồm: chương trình môn học mới, đào tạo lại giảng viên, tích hợp AI vào giáo trình.",
            "Không bao gồm: kỹ năng cá nhân (→ TUONG_LAI.02).",
            "Từ năm học 2024-2025, chúng tôi đưa vào môn học 'Pháp lý và Công nghệ' bắt buộc cho sinh viên năm ba.",
            "NTL_03 (PV03)"
        ),
        (
            "TUONG_LAI.04 – Cạnh tranh thị trường pháp lý",
            "Các đoạn bàn về sự thay đổi cấu trúc cạnh tranh trong thị trường dịch vụ pháp lý do AI tạo ra: công ty luật lớn vs nhỏ, AI startup pháp lý.",
            "Bao gồm: thay đổi mô hình kinh doanh, cạnh tranh giá, startup pháp lý AI.",
            "Không bao gồm: lợi ích của AI trong tổ chức (→ LOI_ICH).",
            "AI đang tạo ra sự tái cấu trúc khá sâu sắc. Một công ty luật vừa như chúng tôi với AI có thể cạnh tranh về tốc độ và giá cả với các hãng luật lớn hơn.",
            "NTL_05 (PV05)"
        ),
        (
            "TUONG_LAI.05 – So sánh Việt Nam với quốc tế",
            "Các đoạn so sánh mức độ phát triển, ứng dụng AI pháp lý của Việt Nam với các nước khác, hoặc bài học từ quốc tế.",
            "Bao gồm: Việt Nam đi sau bao nhiêu năm, kinh nghiệm nước ngoài, bài học từ Mỹ/Anh/Trung Quốc.",
            "Không bao gồm: tuân thủ pháp luật nước ngoài trong hoạt động kinh doanh (→ PHAP_LY.04).",
            "Chúng ta đang đi sau các nước phát triển khoảng 3-5 năm về mặt ứng dụng thực tiễn, và khoảng 5-7 năm về mặt khung pháp lý điều chỉnh.",
            "NTL_03 (PV03)"
        ),
    ]

    for (tieu_de, dinh_nghia, bao_gom, khong_bao_gom, vi_du, nguon) in nodes_g7:
        add_h3(doc, tieu_de, color=(0x40, 0x40, 0x40))
        add_label_value(doc, "Định nghĩa", dinh_nghia)
        add_label_value(doc, "BAO GỒM", bao_gom)
        add_label_value(doc, "KHÔNG BAO GỒM", khong_bao_gom)
        add_body(doc, "Ví dụ trích dẫn:")
        add_quote(doc, vi_du, nguon)
        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================
    # PHẦN 4: BẢNG ÁNH XẠ NODE ↔ PHỎNG VẤN
    # =========================================================
    add_h2(doc, "PHẦN 4: BẢNG ÁNH XẠ NODE THEO TỪNG PHỎNG VẤN")
    add_body(doc,
        "Bảng dưới đây dự đoán trước các node CÓ KHẢ NĂNG XUẤT HIỆN trong từng file phỏng vấn, "
        "giúp người mã hóa tập trung chú ý đúng chỗ. Dấu ● = xuất hiện nhiều/rõ ràng; ○ = có thể xuất hiện; – = ít khả năng."
    )
    doc.add_paragraph()

    mapping_table = doc.add_table(rows=1, cols=6)
    mapping_table.style = 'Table Grid'
    headers_map = ["Node", "PV01\n(Fintech)", "PV02\n(Sản xuất)", "PV03\n(Học giả)", "PV04\n(Nhà nước)", "PV05\n(Luật sư)"]
    hdr_cells = mapping_table.rows[0].cells
    for i, h in enumerate(headers_map):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr_cells[i], "2E74B5")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    mapping_data = [
        ("BCLC.01 – Áp lực công việc",          "●", "●", "○", "●", "●"),
        ("BCLC.02 – Nhu cầu số hóa",             "●", "●", "●", "●", "●"),
        ("BCLC.03 – Lộ trình triển khai",        "●", "●", "–", "●", "○"),
        ("BCLC.04 – Đặc thù ngành",              "●", "●", "○", "●", "○"),
        ("UNG_DUNG.01 – Rà soát hợp đồng",       "●", "●", "–", "–", "●"),
        ("UNG_DUNG.02 – Tra cứu pháp luật",      "●", "○", "–", "●", "●"),
        ("UNG_DUNG.03 – Soạn thảo",              "○", "●", "–", "●", "●"),
        ("UNG_DUNG.04 – Theo dõi tuân thủ",      "●", "●", "–", "●", "○"),
        ("UNG_DUNG.05 – Due diligence M&A",       "–", "–", "–", "–", "●"),
        ("UNG_DUNG.06 – Dịch thuật",             "–", "●", "–", "–", "–"),
        ("UNG_DUNG.07 – Xây dựng VBPL",          "–", "–", "–", "●", "–"),
        ("UNG_DUNG.08 – Tranh chấp/ODR",         "–", "–", "–", "●", "●"),
        ("LOI_ICH.01 – Tiết kiệm thời gian",     "●", "●", "–", "○", "●"),
        ("LOI_ICH.02 – Chất lượng/Chính xác",    "●", "●", "–", "○", "●"),
        ("LOI_ICH.03 – Giảm chi phí",            "○", "●", "–", "–", "●"),
        ("LOI_ICH.04 – Phổ cập pháp lý",         "–", "–", "●", "●", "○"),
        ("THACH_THUC.01 – Dữ liệu TV",           "●", "○", "●", "○", "●"),
        ("THACH_THUC.02 – Bảo mật",              "●", "●", "–", "●", "●"),
        ("THACH_THUC.03 – Ngữ cảnh",             "●", "●", "●", "–", "○"),
        ("THACH_THUC.04 – Kháng cự nhân sự",     "●", "●", "–", "–", "–"),
        ("THACH_THUC.05 – Chi phí/Hạ tầng",      "●", "○", "–", "○", "–"),
        ("PHAP_LY.01 – Khoảng trống pháp lý",    "●", "–", "●", "●", "●"),
        ("PHAP_LY.02 – Trách nhiệm AI",          "●", "–", "●", "–", "○"),
        ("PHAP_LY.03 – Kiến nghị chính sách",    "○", "–", "●", "●", "○"),
        ("PHAP_LY.04 – Đa thẩm quyền",           "–", "●", "–", "●", "–"),
        ("DAO_DUC.01 – AI hỗ trợ/người quyết",   "●", "●", "●", "●", "●"),
        ("DAO_DUC.02 – Thiên vị thuật toán",      "–", "–", "●", "–", "–"),
        ("DAO_DUC.03 – Bảo mật thân chủ",        "–", "–", "●", "–", "●"),
        ("DAO_DUC.04 – Minh bạch KH",            "–", "–", "–", "●", "●"),
        ("TUONG_LAI.01 – Thay đổi vai trò",      "●", "●", "●", "–", "●"),
        ("TUONG_LAI.02 – Kỹ năng AI",            "–", "●", "●", "–", "●"),
        ("TUONG_LAI.03 – Đào tạo/giáo dục",      "–", "●", "●", "–", "–"),
        ("TUONG_LAI.04 – Cạnh tranh TT",         "–", "–", "–", "–", "●"),
        ("TUONG_LAI.05 – So sánh quốc tế",       "–", "●", "●", "●", "–"),
    ]

    for row_data in mapping_data:
        row = mapping_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            if i > 0:
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val == "●":
                set_cell_bg(row[i], "E2EFDA")
            elif val == "○":
                set_cell_bg(row[i], "FFF2CC")

    doc.add_paragraph()
    doc.add_page_break()

    # =========================================================
    # PHẦN 5: HƯỚNG DẪN THAO TÁC TRONG NVIVO
    # =========================================================
    add_h2(doc, "PHẦN 5: HƯỚNG DẪN THAO TÁC CODE SELECTION TRONG NVIVO 14")

    steps = [
        ("Bước 1: Tạo Node Tree trong NVivo",
         [
             "Mở NVivo 14 → Tab Codes → Nhấp phải vào Codes → New Code.",
             "Tạo 7 node chính (parent): NHOM_BOI_CANH, NHOM_UNG_DUNG, NHOM_LOI_ICH, NHOM_THACH_THUC, NHOM_PHAP_LY, NHOM_DAO_DUC, NHOM_TUONG_LAI.",
             "Trong mỗi node chính, tạo các node con theo đúng mã hiệu và tên trong Phần 2.",
             "Nhập Description cho mỗi node bằng cách copy định nghĩa từ Phần 3 vào ô mô tả.",
         ]
        ),
        ("Bước 2: Import file phỏng vấn",
         [
             "Tab Files → Import → chọn 5 file .docx trong thư mục phong_van_AI_phap_che.",
             "Đặt tên file ngắn gọn: PV01, PV02, PV03, PV04, PV05.",
             "Tạo Cases tương ứng và liên kết với Classification 'Người được phỏng vấn' (xem hướng dẫn Case Classification).",
         ]
        ),
        ("Bước 3: Thực hiện Code Selection",
         [
             "Mở file PV01 → Đọc từng câu trả lời của NTL.",
             "Dùng chuột bôi đen (kéo chọn) đoạn văn bản muốn mã hóa.",
             "Nhấp phải → Code Selection → chọn đúng node từ danh sách.",
             "Một đoạn có thể mã hóa vào NHIỀU node: bôi đen → nhấp phải → Code Selection → chọn node thứ nhất → lặp lại với node thứ hai.",
             "Sử dụng phím tắt: Alt + F10 để mở nhanh hộp thoại mã hóa.",
         ]
        ),
        ("Bước 4: Kiểm tra chất lượng mã hóa",
         [
             "Sau khi mã hóa xong PV01, mở từng node để xem tất cả đoạn văn đã được mã hóa vào đó.",
             "Đọc lại và xác nhận các đoạn trong mỗi node đều nhất quán với định nghĩa.",
             "Dùng Query → Text Search để tìm từ khóa liên quan đến từng node còn sót.",
             "Ghi chú những đoạn không chắc vào Memo gắn với file phỏng vấn tương ứng.",
         ]
        ),
        ("Bước 5: Kiểm tra độ tin cậy (Inter-rater Reliability)",
         [
             "Sau khi mã hóa cả 5 file, nhờ người mã hóa thứ hai mã hóa độc lập PV03 (file trung lập nhất – học giả).",
             "So sánh bằng: Explore → Charts → Coding Comparison hoặc tính Kappa thủ công.",
             "Mục tiêu: Cohen's Kappa ≥ 0.70 cho mỗi node.",
             "Thảo luận và điều chỉnh định nghĩa node nếu có bất đồng đáng kể.",
         ]
        ),
    ]

    for (step_title, step_items) in steps:
        add_h3(doc, step_title)
        for item in step_items:
            doc.add_paragraph(item, style='List Number')
        doc.add_paragraph()

    # =========================================================
    # PHẦN 6: GHI CHÚ & CẬP NHẬT
    # =========================================================
    add_h2(doc, "PHẦN 6: NHẬT KÝ PHIÊN BẢN CODEBOOK")

    ver_table = doc.add_table(rows=1, cols=4)
    ver_table.style = 'Table Grid'
    vh = ver_table.rows[0].cells
    for i, h in enumerate(["Phiên bản", "Ngày", "Người chỉnh sửa", "Nội dung thay đổi"]):
        vh[i].text = h
        vh[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(vh[i], "2E74B5")
        vh[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    ver_data = [
        ("v1.0", "Tháng 4/2025", "Lê Thị Hương", "Tạo mới – mã hóa lần đầu từ 5 phỏng vấn"),
        ("v1.x", "...", "...", "Cập nhật sau khi kiểm tra inter-rater reliability"),
    ]
    for row_d in ver_data:
        r = ver_table.add_row().cells
        for i, v in enumerate(row_d):
            r[i].text = v

    doc.add_paragraph()
    add_body(doc, "Lưu ý: Codebook này nên được cập nhật mỗi khi có thay đổi về định nghĩa node, thêm hoặc xóa node, hoặc sau mỗi vòng kiểm tra độ tin cậy. Phiên bản cuối cùng phải được đính kèm vào phụ lục của báo cáo nghiên cứu.")

    doc.save(os.path.join(OUTPUT_DIR, "CODEBOOK_Ma_Hoa_AI_Phap_Che.docx"))
    print("Đã tạo: CODEBOOK_Ma_Hoa_AI_Phap_Che.docx")


if __name__ == "__main__":
    create_codebook()
